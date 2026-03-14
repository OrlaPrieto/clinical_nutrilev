from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import os
import io
import re
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types as genai_types
from google.oauth2 import id_token
from google.auth.transport import requests as google_requests
from dotenv import load_dotenv
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

load_dotenv()

app = Flask(__name__)

# Security: Max upload size is 5MB to prevent OOM
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024

# Security: Rate Limiting to prevent Denial of Wallet on Gemini API
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["200 per day", "10 per minute"],
    storage_uri="memory://"
)

# Security: Restrict CORS to specific origins
CORS(app, resources={r"/api/*": {"origins": ["http://localhost:4200", "https://app.clinicanutrilev.com"]}})

NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# =============================================================================
# DETECCIÓN DE FORMATO
#
# Formato EQUIVALENCIAS (este script):
#   - Menú de equivalentes con 3 columnas (MENÚ 1 / MENÚ 2 / MENÚ 3)
#   - Dish cards: tablas de 2 columnas (C0=imagen vacía, C1=texto)
#   - P0 bold sz15 = "MENÚ 1/2/3"
#   - P1 bold = nombre del platillo
#   - P2 normal = descripción
#   - Secciones: DESAYUNO (T14-16), COMIDA (T20-22), CENA (T26-28)
#     contando desde T0; se detectan dinámicamente por estructura
#
# Formato SEMANAL (script anterior):
#   - Fila 0 = header con días de semana
#   - Col 0 = tiempo de comida
#
# Formato MENU123 (script anterior):
#   - Tabla grande con 5 columnas
#   - Cols 2/3/4 = Menu 1/2/3
# =============================================================================

def detect_format(doc: Document) -> str:
    """
    Detecta si el documento es formato 'equivalencias', 'menu123' o 'semanal'.
    """
    # Equivalencias: tiene dish cards de 2 celdas donde C1 empieza con "MENÚ N"
    for table in doc.tables:
        if len(table.rows) == 1 and len(table.columns) == 2:
            c1_text = table.rows[0].cells[1].text.strip()
            if re.match(r'MEN[ÚU]\s*[123]', c1_text, re.IGNORECASE):
                return "equivalencias"

    # Semanal: header con días de semana
    for table in doc.tables:
        if len(table.rows) >= 3:
            header_text = " ".join(c.text.lower() for c in table.rows[0].cells)
            dias = ["lunes", "martes", "miercoles", "miércoles", "jueves",
                    "viernes", "sabado", "sábado", "domingo"]
            if sum(1 for d in dias if d in header_text) >= 2:
                return "semanal"

    # Menu123: tabla con 5+ columnas y MENÚ 1/2/3 en cabecera
    for table in doc.tables:
        if len(table.columns) >= 5:
            for row in table.rows[:3]:
                for ci in range(2, min(6, len(row.cells))):
                    if re.search(r'(MEN[^A-Z\s]*\s*[123]|DIA\s*[123])',
                                 row.cells[ci].text, re.IGNORECASE):
                        return "menu123"

    return "menu123"


# =============================================================================
# FORMATO EQUIVALENCIAS — IMAGEN SLOTS
#
# El documento tiene dish cards en tablas de 2 columnas:
#   C0 = celda vacía (aquí va la imagen)
#   C1 = texto (MENÚ N + nombre + descripción)
#
# Estructura fija por sección:
#   DESAYUNO: 3 dish cards seguidas de tablas de equivalentes
#   COMIDA:   3 dish cards seguidas de tablas de equivalentes
#   CENA:     3 dish cards seguidas de tablas de equivalentes
#
# Solo DESAYUNO y COMIDA reciben imagen (CENA excluida por instrucción del usuario)
#
# Detección dinámica: recorremos todas las tablas buscando dish cards
# y determinamos la sección según el banner más reciente.
# =============================================================================

SECCIONES_CON_IMAGEN_EQ = {"DESAYUNO", "COMIDA"}


def is_dish_card(table) -> bool:
    """Tabla de 2 columnas, 1 fila, donde C1 empieza con 'MENÚ N'"""
    if len(table.rows) != 1 or len(table.columns) != 2:
        return False
    c1 = table.rows[0].cells[1].text.strip()
    return bool(re.match(r'MEN[ÚU]\s*[123]', c1, re.IGNORECASE))


def is_section_banner(table) -> Optional[str]:
    """
    Tabla de 1 columna con texto de sección (DESAYUNO/COMIDA/CENA).
    Devuelve la sección o None.
    """
    if len(table.rows) != 1 or len(table.columns) != 1:
        return None
    text = table.rows[0].cells[0].text.strip().upper()
    for sec in ["DESAYUNO", "COMIDA", "CENA"]:
        if sec in text and len(text) < 30:
            return sec
    return None


def extract_dish_name_from_card(table) -> str:
    """Extrae el nombre del platillo del párrafo P1 (bold) de la dish card."""
    if len(table.rows) < 1 or len(table.columns) < 2:
        return "platillo"
    cell1 = table.rows[0].cells[1]
    paras = cell1.paragraphs
    # P0 = "MENÚ N", P1 = nombre del platillo (bold)
    if len(paras) >= 2:
        return paras[1].text.strip()
    elif len(paras) >= 1:
        text = paras[0].text.strip()
        # Remove "MENÚ N" prefix
        text = re.sub(r'^MEN[ÚU]\s*[123]\s*', '', text, flags=re.IGNORECASE).strip()
        return text.split('\n')[0].strip()
    return "platillo"


def find_image_slots_equivalencias(doc: Document) -> list:
    """
    Encuentra los slots de imagen para el formato equivalencias.
    Devuelve lista de (table_index, dish_name, has_image, seccion)
    Solo incluye DESAYUNO y COMIDA.
    """
    slots = []
    current_section = None

    for ti, table in enumerate(doc.tables):
        # Actualizar sección actual si encontramos un banner
        sec = is_section_banner(table)
        if sec:
            current_section = sec
            continue

        # Si es dish card y la sección tiene imagen
        if is_dish_card(table) and current_section in SECCIONES_CON_IMAGEN_EQ:
            dish_name = extract_dish_name_from_card(table)
            has_image = bool(
                table.rows[0].cells[0]._tc.findall('.//{%s}blip' % NS_A)
            )
            slots.append((ti, dish_name, has_image, current_section))

    return slots


# =============================================================================
# FORMATO MENU123 y SEMANAL — reutilizados del script anterior
# =============================================================================

def find_menu_table(doc: Document):
    fmt = detect_format(doc)
    for table in doc.tables:
        if fmt == "semanal":
            if len(table.rows) >= 3 and len(table.columns) >= 3:
                col0_texts = [table.rows[i].cells[0].text.lower()
                              for i in range(1, len(table.rows))]
                if any(t in " ".join(col0_texts)
                       for t in ["desayuno", "comida", "cena"]):
                    return table, fmt
        else:
            if len(table.columns) >= 5:
                for row in table.rows[:3]:
                    for ci in [2, 3, 4]:
                        if ci < len(row.cells):
                            if re.search(r'(MEN[^A-Z\s]*\s*[123]|DIA\s*[123])',
                                         row.cells[ci].text, re.IGNORECASE):
                                return table, fmt
    return None, fmt


SECCIONES_CON_IMAGEN = {"DESAYUNO", "COMIDA"}


def is_menu_table_fmt123(table) -> bool:
    if len(table.columns) < 5:
        return False
    for row in table.rows[:3]:
        if any(x in row.cells[0].text.upper()
               for x in ["DESAYUNO", "COMIDA", "CENA"]):
            return True
        for ci in [2, 3, 4]:
            if ci < len(row.cells):
                if re.search(r'(MEN[^A-Z\s]*\s*[123]|\bDIA\s*[123]\b)',
                             row.cells[ci].text, re.IGNORECASE):
                    return True
    return False


def find_image_slots_menu123(doc: Document) -> list:
    slots = []
    for ti, table in enumerate(doc.tables):
        if not is_menu_table_fmt123(table):
            continue
        fmt_sub = "menu13"
        for row in table.rows[:5]:
            for ci in range(2, min(6, len(row.cells))):
                if re.search(r'\bDIA\s*[123]\b', row.cells[ci].text,
                             re.IGNORECASE):
                    fmt_sub = "dia123"

        meal_rows = []
        seccion_actual = "DESAYUNO"
        colaciones_vistas = 0
        prev_was_colacion = False

        for ri, row in enumerate(table.rows):
            if len(row.cells) < 5:
                continue
            col0 = row.cells[0].text.strip()
            if fmt_sub == "dia123":
                col0_up = col0.upper()
                if "DESAYUNO" in col0_up:
                    seccion_actual = "DESAYUNO"
                elif "COMIDA" in col0_up and len(col0_up) < 20:
                    seccion_actual = "COMIDA"
                elif "CENA" in col0_up and len(col0_up) < 20:
                    seccion_actual = "CENA"
            else:
                if "COLACI" in col0.upper():
                    colaciones_vistas += 1
                    prev_was_colacion = True
                    continue
                if prev_was_colacion:
                    prev_was_colacion = False
                    if colaciones_vistas == 1:
                        seccion_actual = "COMIDA"
                    elif colaciones_vistas >= 2:
                        seccion_actual = "CENA"

            meal_names_in_row = {}
            for ci in [2, 3, 4]:
                txt = row.cells[ci].text.strip()
                if re.search(r'MEN[^A-Z\s]*\s*[123]', txt, re.IGNORECASE) or \
                   (fmt_sub == "dia123" and re.search(
                       r'\bDIA\s*[123]\b', txt, re.IGNORECASE)):
                    name = re.sub(
                        r'^(MEN[^A-Z\s]*\s*[123]|\bDIA\s*[123]\b)\s*', '',
                        txt, flags=re.IGNORECASE).strip()
                    meal_names_in_row[ci] = name
            if meal_names_in_row:
                meal_rows.append((ri, meal_names_in_row, seccion_actual))

        for idx, (meal_ri, names, seccion) in enumerate(meal_rows):
            next_ri = (meal_rows[idx + 1][0] if idx + 1 < len(meal_rows)
                       else len(table.rows))
            image_row = None
            for search_ri in range(next_ri - 1, meal_ri, -1):
                if search_ri >= len(table.rows):
                    continue
                row = table.rows[search_ri]
                has_content = any(
                    row.cells[ci].text.strip() and
                    "COLACI" not in row.cells[ci].text.upper() and
                    len(row.cells[ci].text.strip()) > 3
                    for ci in [2, 3, 4] if ci < len(row.cells)
                )
                if has_content:
                    is_name_row = any(
                        re.search(
                            r'(MEN[^A-Z\s]*\s*[123]|\bDIA\s*[123]\b)',
                            row.cells[ci].text, re.IGNORECASE)
                        for ci in [2, 3, 4] if ci < len(row.cells)
                    )
                    if not is_name_row:
                        image_row = search_ri
                        break

            if image_row is not None and seccion in SECCIONES_CON_IMAGEN:
                for ci in [2, 3, 4]:
                    if ci >= len(table.rows[image_row].cells):
                        continue
                    cell = table.rows[image_row].cells[ci]
                    has_image = bool(
                        cell._tc.findall('.//{%s}blip' % NS_A))
                    meal_name = names.get(ci, f"platillo col{ci}")
                    slots.append(
                        (ti, image_row, ci, meal_name, has_image, seccion))

    return slots


TIEMPOS_CON_IMAGEN_SEMANAL = {"desayuno", "comida"}


def find_image_slots_semanal(doc: Document) -> list:
    slots = []
    table, fmt = find_menu_table(doc)
    if not table or fmt != "semanal":
        return slots
    ti = next((i for i, t in enumerate(doc.tables) if t is table), 0)
    n_cols = len(table.rows[0].cells)
    for ri, row in enumerate(table.rows[1:], start=1):
        tiempo = row.cells[0].text.strip().lower()
        if not any(t in tiempo for t in TIEMPOS_CON_IMAGEN_SEMANAL):
            continue
        for ci in range(1, n_cols):
            if ci >= len(row.cells):
                continue
            cell = row.cells[ci]
            txt = cell.text.strip()
            if not txt:
                continue
            meal_name = (txt.split(":")[0].strip() if ":" in txt
                         else " ".join(txt.split()[:4]).strip())
            has_image = bool(cell._tc.findall('.//{%s}blip' % NS_A))
            slots.append((ti, ri, ci, meal_name, has_image, tiempo.upper()))
    return slots


# =============================================================================
# INSERCIÓN DE IMAGEN
# =============================================================================

def insert_image_in_cell_equivalencias(doc: Document, table_index: int,
                                        image_bytes: bytes):
    """
    Inserta imagen en C0 de la dish card (tabla de 2 col, 1 fila).
    C0 es la celda izquierda reservada para imagen.
    """
    table = doc.tables[table_index]
    cell = table.rows[0].cells[0]

    # Limpiar dibujos existentes
    for drw in cell._tc.findall('.//{%s}drawing' % NS_W):
        drw.getparent().remove(drw)

    # Insertar nueva imagen centrada
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    # Limpiar runs anteriores
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Inches(1.3))


def insert_image_in_cell(doc: Document, ti: int, ri: int, ci: int,
                          image_bytes: bytes, fmt: str = "menu123"):
    """Inserción para formatos menu123 y semanal (lógica original)."""
    table = doc.tables[ti]
    cell = table.rows[ri].cells[ci]
    existing_text = cell.text.strip()

    for drw in cell._tc.findall('.//{%s}drawing' % NS_W):
        drw.getparent().remove(drw)

    if fmt == "semanal":
        for p in list(cell.paragraphs):
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
        img_para = cell.paragraphs[0]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_para.add_run()
        img_run.add_picture(io.BytesIO(image_bytes), width=Inches(0.85))
        if existing_text:
            txt_para = cell.add_paragraph()
            txt_para.add_run(existing_text)
    else:
        para = cell.paragraphs[0]
        for r in list(para.runs):
            r._r.getparent().remove(r._r)
        run = para.add_run()
        run.add_picture(io.BytesIO(image_bytes), width=Inches(1.0))
        if existing_text:
            para.add_run("\n" + existing_text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# =============================================================================
# GENERACIÓN DE IMÁGENES (Imagen 4 → Imagen 3 fallback)
# =============================================================================

def generate_image_gemini(meal_name: str, gemini_key: str) -> Optional[bytes]:
    prompt = (
        f"Professional food photography of '{meal_name}', healthy Mexican cuisine, "
        "served on a white plate, overhead or 45-degree angle, natural lighting, "
        "clean background, appetizing, no text, no watermarks"
    )
    client = genai.Client(api_key=gemini_key)
    for model_name in ["imagen-4.0-generate-001", "imagen-4.0-fast-generate-001",
                        "imagen-3.0-generate-002"]:
        try:
            response = client.models.generate_images(
                model=model_name,
                prompt=prompt,
                config=genai_types.GenerateImagesConfig(
                    number_of_images=1,
                    aspect_ratio="1:1",
                    safety_filter_level="block_low_and_above",
                ),
            )
            return response.generated_images[0].image.image_bytes
        except Exception as e:
            if any(x in str(e).upper() for x in ["PAID", "BILLING", "QUOTA"]):
                print(f"[Imagen] Sin acceso para {model_name}, probando siguiente...")
            else:
                print(f"[Imagen] Error con {model_name}: {e}")
            continue
    return None


# =============================================================================
# EXTRACCIÓN DE DATOS DEL MENÚ (para lista de compras)
# =============================================================================

def extract_equivalencias_data(doc: Document) -> dict:
    """
    Extrae todos los textos del menú de equivalencias para generar
    la lista de compras. Recoge:
    - Nombre y descripción de cada dish card
    - Contenido de cada fila de equivalentes (cols M1/M2/M3)
    - Texto de colaciones
    """
    menus = {
        "Menu 1": {"dias": 2, "comidas": {}, "cols": [2]},
        "Menu 2": {"dias": 2, "comidas": {}, "cols": [3]},
        "Menu 3": {"dias": 3, "comidas": {}, "cols": [4]},
    }
    current_section = None
    all_text_blocks = []

    for ti, table in enumerate(doc.tables):
        # Banner de sección
        sec = is_section_banner(table)
        if sec:
            current_section = sec
            continue

        # Dish card: extraer nombre + descripción
        if is_dish_card(table):
            cell1 = table.rows[0].cells[1]
            paras = cell1.paragraphs
            menu_label = ""
            dish_name = ""
            description = ""
            for pi, para in enumerate(paras):
                text = para.text.strip()
                if pi == 0:
                    menu_label = text  # "MENÚ 1"
                elif pi == 1:
                    dish_name = text
                elif pi == 2:
                    description = text

            # Determinar a qué menú pertenece
            m_num = re.search(r'[123]', menu_label)
            if m_num and current_section:
                label = f"Menu {m_num.group()}"
                sec_key = current_section
                if label in menus:
                    existing = menus[label]["comidas"].get(sec_key, "")
                    menus[label]["comidas"][sec_key] = (
                        existing + f"\n{dish_name}: {description}"
                        if existing else f"{dish_name}: {description}"
                    )
            all_text_blocks.append(f"{dish_name} - {description}")
            continue

        # Colaciones (banner de 1 col con "COLACIÓN")
        if (len(table.rows) == 1 and len(table.columns) == 1):
            text = table.rows[0].cells[0].text.strip()
            if "COLACI" in text.upper():
                all_text_blocks.append(text)
                # Agregar a todos los menús
                for label in menus:
                    sec_key = "COLACION_" + (
                        "MAT" if "MATUTINA" in text.upper() else "VES")
                    menus[label]["comidas"][sec_key] = text
                continue

        # Tabla de equivalentes (5 cols: cant | tipo | M1 | M2 | M3)
        if len(table.columns) == 5 and current_section:
            for row in table.rows:
                cells = row.cells
                if len(cells) < 5:
                    continue
                # Saltar fila de encabezado (MENÚ 1/2/3)
                if re.search(r'MEN[ÚU]\s*[123]', cells[2].text, re.IGNORECASE):
                    continue
                # Saltar nota (span)
                if cells[0].text.strip() == "" and cells[1].text.strip() == "":
                    nota = cells[2].text.strip()
                    if nota:
                        all_text_blocks.append(nota)
                    continue

                amount = cells[0].text.strip()
                tipo = cells[1].text.strip()
                for ci, label in [(2, "Menu 1"), (3, "Menu 2"), (4, "Menu 3")]:
                    txt = cells[ci].text.strip()
                    if txt and label in menus:
                        existing = menus[label]["comidas"].get(
                            current_section, "")
                        entry = (f"{amount} equiv {tipo}: {txt}" if amount
                                 else txt)
                        menus[label]["comidas"][current_section] = (
                            existing + f"\n{entry}" if existing else entry
                        )
                        all_text_blocks.append(txt)

    return {"menus": menus, "todos_ingredientes": all_text_blocks}


def extract_menu_data(doc: Document) -> dict:
    fmt = detect_format(doc)
    if fmt == "equivalencias":
        return extract_equivalencias_data(doc)
    elif fmt == "semanal":
        return extract_semanal(doc)
    else:
        return extract_menu123(doc)


def extract_semanal(doc: Document) -> dict:
    table, _ = find_menu_table(doc)
    if not table:
        return {"menus": {}, "todos_ingredientes": []}
    header_row = table.rows[0]
    n_cols = len(header_row.cells)
    col_multiplicadores = {}
    menu_counter = 1
    menus = {}
    menu_3_cols = []
    assigned_menu = {}
    for ci in range(1, n_cols):
        header = header_row.cells[ci].text.strip()
        mult = (len([x for x in header.split("/") if x.strip()])
                if "/" in header else 1)
        col_multiplicadores[ci] = mult
    for ci in range(1, n_cols):
        if col_multiplicadores[ci] >= 2:
            label = f"Menu {menu_counter}"
            menus[label] = {"dias": col_multiplicadores[ci],
                            "comidas": {}, "cols": [ci]}
            assigned_menu[ci] = label
            menu_counter += 1
        else:
            menu_3_cols.append(ci)
    for i, ci in enumerate(menu_3_cols):
        label = f"Menu 3{'abc'[i] if i < 3 else str(i)}"
        menus[label] = {"dias": 1, "comidas": {}, "cols": [ci]}
        assigned_menu[ci] = label
    for ri in range(1, len(table.rows)):
        row = table.rows[ri]
        tiempo = row.cells[0].text.strip()
        if not tiempo:
            continue
        for ci in range(1, n_cols):
            if ci >= len(row.cells):
                continue
            menu_label = assigned_menu.get(ci)
            if menu_label and menu_label in menus:
                contenido = row.cells[ci].text.strip()
                if contenido:
                    menus[menu_label]["comidas"][tiempo] = contenido
    todos = [txt for m in menus.values() for txt in m["comidas"].values()]
    return {"menus": menus, "todos_ingredientes": todos}


def extract_menu123(doc: Document) -> dict:
    table, _ = find_menu_table(doc)
    if not table:
        return {"menus": {}, "todos_ingredientes": []}
    fmt_sub = "menu13"
    for row in table.rows[:5]:
        for ci in range(2, min(6, len(row.cells))):
            if re.search(r'\bDIA\s*[123]\b', row.cells[ci].text, re.IGNORECASE):
                fmt_sub = "dia123"
                break
    menus = {
        "Menu 1": {"dias": 2, "comidas": {}, "cols": [2]},
        "Menu 2": {"dias": 2, "comidas": {}, "cols": [3]},
        "Menu 3": {"dias": 3, "comidas": {}, "cols": [4]},
    }
    seccion_actual = "DESAYUNO"
    colaciones_vistas = 0
    prev_was_colacion = False
    for row in table.rows:
        if len(row.cells) < 5:
            continue
        col0 = row.cells[0].text.strip()
        if fmt_sub == "dia123":
            col0_up = col0.upper()
            if "DESAYUNO" in col0_up:
                seccion_actual = "DESAYUNO"
            elif "COMIDA" in col0_up and len(col0_up) < 20:
                seccion_actual = "COMIDA"
            elif "CENA" in col0_up and len(col0_up) < 20:
                seccion_actual = "CENA"
        else:
            if "COLACI" in col0.upper():
                colaciones_vistas += 1
                prev_was_colacion = True
                continue
            if prev_was_colacion:
                prev_was_colacion = False
                if colaciones_vistas == 1:
                    seccion_actual = "COMIDA"
                elif colaciones_vistas >= 2:
                    seccion_actual = "CENA"
        for ci, menu_label in [(2, "Menu 1"), (3, "Menu 2"), (4, "Menu 3")]:
            if ci >= len(row.cells):
                continue
            cell_text = row.cells[ci].text.strip()
            if not cell_text:
                continue
            is_meal_name = bool(re.search(
                r'(MEN[^A-Z\s]*\s*[123]|\bDIA\s*[123]\b)',
                cell_text, re.IGNORECASE))
            if is_meal_name:
                nombre = re.sub(
                    r'^(MEN[^A-Z\s]*\s*[123]|\bDIA\s*[123]\b)\s*', '',
                    cell_text, flags=re.IGNORECASE).strip()
                menus[menu_label]["comidas"][seccion_actual] = (
                    menus[menu_label]["comidas"].get(seccion_actual, "")
                    + "\n" + nombre)
            else:
                menus[menu_label]["comidas"][seccion_actual] = (
                    menus[menu_label]["comidas"].get(seccion_actual, "")
                    + "\n" + cell_text)
    todos = [txt for m in menus.values() for txt in m["comidas"].values()]
    return {"menus": menus, "todos_ingredientes": todos}


# =============================================================================
# PROMPT DE LISTA DE COMPRAS
# =============================================================================

def build_shopping_prompt(menu_data: dict) -> str:
    menus = menu_data["menus"]
    bloques = []
    for label, info in menus.items():
        dias = info["dias"]
        bloques.append(f"\n{'='*40}")
        bloques.append(
            f"{label.upper()} — se repite {dias} "
            f"{'día' if dias == 1 else 'días'} a la semana")
        bloques.append(f"{'='*40}")
        for tiempo, contenido in info["comidas"].items():
            bloques.append(f"\n{tiempo}:\n{contenido.strip()}")

    menu_txt = "\n".join(bloques)
    menu_1_dias = next(
        (v["dias"] for k, v in menus.items() if "1" in k), 2)
    menu_2_dias = next(
        (v["dias"] for k, v in menus.items() if "2" in k), 2)

    return f"""Eres una nutrióloga mexicana experta en planificación alimentaria semanal.
Analiza el siguiente plan de menú de 7 días y genera una lista de compras COMPLETA.

{menu_txt}

REGLAS DE CANTIDAD:
- Menú 1 se consume {menu_1_dias} días → multiplica TODAS sus porciones x{menu_1_dias}
- Menú 2 se consume {menu_2_dias} días → multiplica TODAS sus porciones x{menu_2_dias}
- Menú 3 se consume 3 días → suma ingredientes de esos 3 días
- Si un ingrediente aparece en VARIOS menús, SUMA todas las cantidades de los 7 días
- Convierte a unidades de compra prácticas: piezas, gramos, litros, paquetes, bolsas

FORMATO DE SALIDA — tabla Markdown por grupo de alimentos:

## [EMOJI_GRUPO] NOMBRE DEL GRUPO
| Icono | Alimento | Cantidad total 7 días | 💡 Tip de compra | ⭐ Marca recomendada |
|-------|----------|----------------------|------------------|---------------------|
| 🥚 | Huevo entero | 8 piezas | Revisar que no tengan grietas | Bachoco / San Juan |

GRUPOS REQUERIDOS (en este orden):
## 🥩 PROTEÍNAS (carnes, pollo, pescado, atún, sardina, huevo, quesos proteicos)
## 🧀 LÁCTEOS (leche, yogurt, quesos frescos)
## 🌾 CEREALES Y CARBOHIDRATOS (pan, tortillas, arroz, avena, cereal, papa, tostadas, galletas)
## 🫘 LEGUMINOSAS (frijoles, garbanzo, lentejas, hummus)
## 🥑 GRASAS SALUDABLES (aguacate, nuez, almendra, semillas, aceite, crema de cacahuate, mantequilla)
## 🥦 VERDURAS Y HORTALIZAS (todas las verduras, hongos, nopales, germinados)
## 🍓 FRUTAS (todas las frutas frescas y secas)
## 🧴 CONDIMENTOS Y OTROS (especias, endulzantes, bebidas, suplementos, aderezos)

REGLAS IMPORTANTES:
1. NO omitas ningún alimento aunque aparezca una sola vez en el menú
2. Consolida el mismo alimento de diferentes tiempos sumando cantidades
3. Usa emojis de WhatsApp representativos para cada alimento en la columna Icono
4. El Tip debe ser práctico y específico para supermercados en México
5. La Marca debe ser una marca real disponible en México (Walmart, Soriana, Chedraui)
6. Responde ÚNICAMENTE con las tablas, sin texto antes ni después
"""


def generate_shopping_list(menu_data: dict, gemini_key: str) -> str:
    try:
        client = genai.Client(api_key=gemini_key)
        response = client.models.generate_content(
            model="gemini-flash-latest",
            contents=build_shopping_prompt(menu_data),
        )
        return response.text.strip()
    except Exception as e:
        print(f"[Gemini] Error lista de compras: {e}")
        items = "\n".join(
            f"- {i[:80]}" for i in menu_data["todos_ingredientes"][:50])
        return f"LISTA DE COMPRAS\n\n{items}"


# =============================================================================
# ESCRITURA DE LISTA EN EL DOCUMENTO — estilo del menú de equivalencias
#
# El documento tiene:
#   T30 = Banner "🛒 LISTA DE COMPRAS..." (1 col)
#   T31-T36 = Tablas de compras por grupo (ya existen en menu_v2)
#   T37 = Banner "✅ RECOMENDACIONES"
#   T38 = Tabla 2 col de recomendaciones
#
# Estrategia: eliminar las tablas de compras existentes (T31-T36 o equivalentes)
# y reemplazarlas con las nuevas generadas por Gemini, manteniendo el mismo
# estilo visual (colores por grupo, 5 columnas).
# =============================================================================

# Colores por grupo (header dark, row alt light, tip yellow)
GROUP_COLORS = {
    "PROTEÍNAS":    {"dark": "1A5276", "light": "D4E6F1"},
    "PROTEINAS":    {"dark": "1A5276", "light": "D4E6F1"},
    "LÁCTEOS":      {"dark": "5B2D8E", "light": "F0E6FF"},
    "LACTEOS":      {"dark": "5B2D8E", "light": "F0E6FF"},
    "CEREALES":     {"dark": "856404", "light": "FFF3CD"},
    "CARBOHIDRATOS":{"dark": "856404", "light": "FFF3CD"},
    "LEGUMINOSAS":  {"dark": "2D6A4F", "light": "C7EFCF"},
    "GRASAS":       {"dark": "7D5A00", "light": "F9E4B7"},
    "VERDURAS":     {"dark": "2D6A4F", "light": "C7EFCF"},
    "HORTALIZAS":   {"dark": "2D6A4F", "light": "C7EFCF"},
    "FRUTAS":       {"dark": "A0522D", "light": "FFE5B4"},
    "CONDIMENTOS":  {"dark": "444444", "light": "F0F0F0"},
    "OTROS":        {"dark": "444444", "light": "F0F0F0"},
}
TIP_COLOR = "FFF9E6"
COL_WIDTHS = [396, 2069, 1701, 4526, 2108]  # DXA, sum=10800


def _get_group_colors(group_title: str) -> dict:
    title_up = group_title.upper()
    for key, colors in GROUP_COLORS.items():
        if key in title_up:
            return colors
    return {"dark": "444444", "light": "F0F0F0"}


def _set_cell_shading(cell, fill_hex: str):
    """Aplica color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    # Remove existing shd if any
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "CCCCCC", size: int = 4):
    """Aplica bordes a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def _set_col_width(cell, width_dxa: int):
    """Establece el ancho de una celda en DXA."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcPr.append(tcW)


def _add_shopping_group_table(doc: Document, group_title: str,
                               rows_data: list):
    """
    Agrega una tabla de compras con el estilo del documento de referencia:
      Row 0: header de grupo (span 5 cols, fondo dark)
      Row 1: encabezados de columna (fondo light)
      Rows 2+: datos alternando white/light, tip en amarillo
    """
    colors = _get_group_colors(group_title)
    dark = colors["dark"]
    light = colors["light"]

    # Tabla: header grupo + col headers + data rows
    n_data = len(rows_data)
    tbl = doc.add_table(rows=2 + n_data, cols=5)

    # --- Row 0: header de grupo (span completo) ---
    row0 = tbl.rows[0]
    # Fusionar las 5 celdas
    from docx.oxml import OxmlElement as OE
    cell0 = row0.cells[0]
    # Merge across all columns
    cell0.merge(row0.cells[4])
    _set_cell_shading(cell0, dark)
    _set_cell_borders(cell0, dark)
    p0 = cell0.paragraphs[0]
    run0 = p0.add_run(group_title)
    run0.bold = True
    run0.font.size = Pt(11)
    run0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # --- Row 1: encabezados de columna ---
    row1 = tbl.rows[1]
    headers = ["", "Alimento", "Cantidad 7 días",
               "💡 Tip de compra", "⭐ Marca recomendada"]
    for ci, (hdr, w) in enumerate(zip(headers, COL_WIDTHS)):
        cell = row1.cells[ci]
        _set_cell_shading(cell, light)
        _set_cell_borders(cell, dark, 4)
        _set_col_width(cell, w)
        p = cell.paragraphs[0]
        r = p.add_run(hdr)
        r.bold = True
        r.font.size = Pt(9)

    # --- Rows 2+: datos ---
    for di, row_data in enumerate(rows_data):
        row = tbl.rows[2 + di]
        row_fill = "FFFFFF" if di % 2 == 0 else light
        for ci, (val, w) in enumerate(zip(row_data[:5], COL_WIDTHS)):
            cell = row.cells[ci]
            fill = TIP_COLOR if ci == 3 else row_fill
            _set_cell_shading(cell, fill)
            _set_cell_borders(cell, "CCCCCC", 2)
            _set_col_width(cell, w)
            p = cell.paragraphs[0]
            r = p.add_run(str(val) if val else "")
            r.font.size = Pt(9)
            if ci == 1:  # Alimento en negrita
                r.bold = True


def _parse_shopping_markdown(text: str) -> list:
    """
    Parsea el markdown de Gemini en:
    [(titulo_grupo, [[icono, alimento, cant, tip, marca], ...]), ...]
    """
    groups = []
    current_title = None
    current_rows = []

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            if current_title is not None:
                groups.append((current_title, current_rows))
            current_title = line[3:].strip()
            current_rows = []
            continue
        if re.match(r'^\|[\s\-\|:]+\|$', line):
            continue
        if line.startswith("|"):
            parts = [c.strip() for c in line.split("|") if c.strip()]
            if not parts:
                continue
            if parts[0].lower() in ("icono", "icon", "emoji"):
                continue
            while len(parts) < 5:
                parts.append("")
            current_rows.append(parts[:5])

    if current_title is not None:
        groups.append((current_title, current_rows))

    return groups


def _find_shopping_section_tables(doc: Document) -> tuple:
    """
    Encuentra el índice de tabla donde empieza la sección de compras
    (banner '🛒 LISTA DE COMPRAS') y dónde termina (banner 'RECOMENDACIONES').
    Devuelve (start_idx, end_idx) de las tablas a reemplazar.
    start_idx = índice del banner de compras
    end_idx = índice del banner de recomendaciones (exclusivo)
    """
    start_idx = None
    end_idx = None

    for ti, table in enumerate(doc.tables):
        if len(table.rows) == 1 and len(table.columns) == 1:
            text = table.rows[0].cells[0].text.strip().upper()
            if "LISTA DE COMPRAS" in text:
                start_idx = ti
            elif "RECOMENDACIONES" in text and start_idx is not None:
                end_idx = ti
                break

    return start_idx, end_idx


def replace_shopping_tables(doc: Document, markdown_text: str):
    """
    Reemplaza las tablas de compras existentes entre el banner de LISTA DE COMPRAS
    y el banner de RECOMENDACIONES con las nuevas generadas por Gemini.
    
    Si no existen tablas de compras previas, las inserta después del banner.
    Mantiene el banner de LISTA DE COMPRAS y el de RECOMENDACIONES intactos.
    """
    start_idx, end_idx = _find_shopping_section_tables(doc)

    if start_idx is None:
        # No hay sección de compras, agregar al final
        doc.add_page_break()
        p = doc.add_paragraph()
        r = p.add_run("🛒  LISTA DE COMPRAS SEMANAL (7 DÍAS)")
        r.bold = True
        r.font.size = Pt(14)
        groups = _parse_shopping_markdown(markdown_text)
        for group_title, rows_data in groups:
            if rows_data:
                _add_shopping_group_table(doc, group_title, rows_data)
                doc.add_paragraph()
        return

    # Hay sección existente: eliminar tablas entre start+1 y end_idx
    # (el banner de compras en start_idx se mantiene)
    if end_idx is None:
        end_idx = len(doc.tables)

    # Tablas a eliminar: las que están entre start_idx+1 y end_idx-1
    tables_to_remove = doc.tables[start_idx + 1:end_idx]

    # Obtener el elemento XML del banner de compras para insertar después de él
    banner_table_el = doc.tables[start_idx]._tbl

    # Remover tablas y párrafos intermedios
    # Las tablas en python-docx están en el body como elementos <w:tbl>
    # Necesitamos también remover los <w:p> vacíos entre ellas
    body = doc.element.body

    # Collect elements to remove (tables between start+1 and end_idx)
    tables_xml = [t._tbl for t in tables_to_remove]

    # Also collect paragraphs between banner and recomendaciones banner
    recom_table_el = (doc.tables[end_idx]._tbl
                      if end_idx < len(doc.tables) else None)

    elements_to_remove = []
    found_banner = False
    found_recom = False
    for child in list(body):
        if child is banner_table_el:
            found_banner = True
            continue
        if recom_table_el is not None and child is recom_table_el:
            found_recom = True
            break
        if found_banner and not found_recom:
            elements_to_remove.append(child)

    for el in elements_to_remove:
        body.remove(el)

    # Ahora insertar las nuevas tablas después del banner de compras
    # Encontrar posición del banner en el body
    banner_pos = list(body).index(banner_table_el)

    groups = _parse_shopping_markdown(markdown_text)

    insert_pos = banner_pos + 1
    for group_title, rows_data in groups:
        if not rows_data:
            continue

        # Agregar párrafo espaciador
        sep_p = OxmlElement('w:p')
        body.insert(insert_pos, sep_p)
        insert_pos += 1

        # Crear tabla temporalmente al final del doc y luego moverla
        _add_shopping_group_table(doc, group_title, rows_data)
        # La tabla recién creada es la última en doc.tables
        new_tbl_el = doc.tables[-1]._tbl
        # Mover al lugar correcto
        body.remove(new_tbl_el)
        body.insert(insert_pos, new_tbl_el)
        insert_pos += 1


# Alias para compatibilidad con el flujo principal
def append_shopping_list(doc: Document, markdown_text: str):
    replace_shopping_tables(doc, markdown_text)


# =============================================================================
# RUTAS API
# =============================================================================

AUTHORIZED_EMAILS = [
    'orla08i@gmail.com',
    'velvetdelacruzvillegas@gmail.com'
]

from functools import wraps
from flask import request, jsonify

def require_google_auth(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # En development local o producción, chequear header
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({'error': 'Falta el token de autorización (Bearer Token)'}), 401
        
        token = auth_header.split(' ')[1]
        try:
            # Especificar el Request object de google-auth
            # En un entorno real, deberías pasar el CLIENT_ID de tu login de Google a id_token.verify_oauth2_token()
            # Como puede depender del env, validaremos que fue emitido por Google y la firma es válida:
            idinfo = id_token.verify_oauth2_token(token, google_requests.Request())
            
            # Verificar si el email está autorizado
            email = idinfo.get('email', '').lower()
            if email not in AUTHORIZED_EMAILS:
                return jsonify({'error': f'Acceso denegado para {email}'}), 403
            
            # Guardamos info en request context por si se necesitara usar después
            request.user_email = email
            
        except ValueError as e:
            # Invalid token
            return jsonify({'error': 'Token inválido', 'details': str(e)}), 401
            
        return f(*args, **kwargs)
    return decorated_function


@app.route('/api/process-menu', methods=['POST'])
@app.route('/process-menu', methods=['POST'])
@require_google_auth
def process_menu():
    if 'file' not in request.files:
        return jsonify({"error": "No se recibió ningún archivo"}), 400

    file = request.files['file']
    gemini_key = request.form.get('api_key') or os.getenv('GEMINI_API_KEY')

    if not gemini_key:
        return jsonify({"error": "Falta la API key de Gemini"}), 400

    try:
        doc = Document(io.BytesIO(file.read()))
        fmt = detect_format(doc)
        print(f"[Formato] {fmt}")

        # ── 1. IMÁGENES ────────────────────────────────────────────────────
        try:
            if fmt == "equivalencias":
                slots = find_image_slots_equivalencias(doc)
                empty = [(ti, name, sec)
                         for ti, name, has, sec in slots if not has]
                print(f"[Imágenes] {len(slots)} slots, {len(empty)} sin imagen")
                for ti, meal_name, sec in empty:
                    print(f"  [{sec}] Generando: {meal_name}")
                    img_bytes = generate_image_gemini(meal_name, gemini_key)
                    if img_bytes:
                        insert_image_in_cell_equivalencias(doc, ti, img_bytes)
                        print(f"  ✓ {meal_name}")
                    else:
                        print(f"  ✗ {meal_name} — sin imagen")
            else:
                if fmt == "semanal":
                    slots = find_image_slots_semanal(doc)
                else:
                    slots = find_image_slots_menu123(doc)
                empty = [(ti, ri, ci, name, sec)
                         for ti, ri, ci, name, has, sec in slots if not has]
                print(f"[Imágenes] {len(slots)} slots, {len(empty)} sin imagen")
                for ti, ri, ci, meal_name, sec in empty:
                    print(f"  [{sec}] Generando: {meal_name}")
                    img_bytes = generate_image_gemini(meal_name, gemini_key)
                    if img_bytes:
                        insert_image_in_cell(doc, ti, ri, ci, img_bytes, fmt)
                        print(f"  ✓ {meal_name}")
                    else:
                        print(f"  ✗ {meal_name} — sin imagen")
        except Exception as e:
            import traceback
            print(f"[Imágenes] Error omitido: {e}")
            print(traceback.format_exc())

        # ── 2. LISTA DE COMPRAS ───────────────────────────────────────────
        menu_data = extract_menu_data(doc)
        shopping_text = generate_shopping_list(menu_data, gemini_key)
        append_shopping_list(doc, shopping_text)

        # ── 3. GUARDAR Y DEVOLVER ─────────────────────────────────────────
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)

        return send_file(
            out,
            as_attachment=True,
            download_name=f"menu_procesado_{file.filename}",
            mimetype=(
                'application/vnd.openxmlformats-officedocument'
                '.wordprocessingml.document')
        )

    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return jsonify({"error": str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({"status": "healthy"}), 200


@app.route('/api', methods=['GET'])
@app.route('/', methods=['GET'])
def root():
    return jsonify({"message": "Nutrilev API is running"}), 200


if __name__ == '__main__':
    app.run(debug=True)