"""
ai_service.py — Nutrilev AI Menu Orchestrator v4.0
====================================================
Orquestador principal que integra Gemini y la generación de DOCX.
"""

from services.docx_utils import (
    Document, _build_docx, GRUPOS_PERMITIDOS, _norm_g
)
from services.gemini_engine import _call_gemini, _resolve_model

def extract_retry_delay(e: Exception) -> float:
    import re
    err_str = str(e)
    
    # 1. Match 'Please retry in X.XXs'
    match = re.search(r"retry\s+in\s+(\d+(?:\.\d+)?)s", err_str, re.IGNORECASE)
    if match:
        try:
            return float(match.group(1))
        except ValueError:
            pass
            
    # 2. Match 'retryDelay': 'Xs'
    match = re.search(r"'retryDelay':\s*'(\d+(?:\.\d+)?)s'", err_str)
    if match:
        try:
            return float(match.group(1))
        except ValueError:
            pass
            
    return 0.0

def _normalizar_menu(menu_json: dict) -> dict:
    """Garantiza que los 3 menús tengan las mismas equivalencias que el esquema_dia."""
    esquema = menu_json.get("esquema_dia", {})
    menus   = menu_json.get("menus", menu_json)
    MK      = ["menu_1", "menu_2", "menu_3"]

    tiempo_map = {
        "desayuno":          ("platillo_solido",  "desayuno_platillo"),
        "desayuno_licuado":  (None,               "desayuno_licuado"),
        "comida":            (None,               "comida"),
        "cena":              (None,               "cena"),
        "colacion_matutina": (None,               "colacion_matutina"),
        "colacion_vespertina": (None,             "colacion_vespertina"),
    }

    for mk in MK:
        menu = menus.get(mk, {})
        for tiempo_key, (sub_key, esquema_key) in tiempo_map.items():
            eq_fijas = esquema.get(esquema_key, {}).get("equivalencias_fijas", [])
            if not eq_fijas: continue

            if tiempo_key == "desayuno" and sub_key:
                bloque = menu.get("desayuno", {}).get("platillo_solido", {})
            elif tiempo_key == "desayuno_licuado":
                bloque = menu.get("desayuno", {}).get("licuado", {})
            else:
                bloque = menu.get(tiempo_key, {})

            if not bloque: continue

            desc_map = { _norm_g(eq.get("grupo", "")): eq.get("descripcion", "") 
                         for eq in bloque.get("equivalencias", []) 
                         if _norm_g(eq.get("grupo", "")) in GRUPOS_PERMITIDOS }

            nuevas_eq = []
            for eq_f in eq_fijas:
                g = eq_f.get("grupo", ""); g_flex = _norm_g(g)
                if g_flex not in GRUPOS_PERMITIDOS: continue
                desc = desc_map.get(g_flex) or desc_map.get(g) or (bloque.get("descripcion", g) if not bloque.get("equivalencias") else g)
                nuevas_eq.append({ "grupo": g, "porciones": eq_f.get("porciones", 0), "descripcion": desc })
            
            if nuevas_eq: bloque["equivalencias"] = nuevas_eq
    return menu_json

def get_base_menu_text() -> str:
    try:
        import os
        base_path = "/Users/orla09i/Desktop/Projects/clinical_nutrilev/apps/api-python/utils/menu_examples/menu edited.docx"
        if not os.path.exists(base_path): return "Standard Nutrilev Structure"
        doc = Document(base_path)
        text = "\n".join([p.text for p in doc.paragraphs] + [c.text for t in doc.tables for r in t.rows for c in r.cells])
        return text[:3000]
    except Exception: return "Standard Nutrilev Structure"

def generate_full_menu_docx(historial_paciente, calorias_objetivo, notas_personalizadas, menu_base_texto, gemini_key):
    print(f"[NutriArchitect v4.0] Generando para: {historial_paciente.get('nombre', '?')}")
    
    # 1. Llamar a Gemini
    menu_json = _call_gemini(historial_paciente, calorias_objetivo, notas_personalizadas, menu_base_texto, gemini_key)
    if not menu_json:
        raise ValueError("No se pudo generar el menú con IA.")

    # 2. Normalizar
    menu_json = _normalizar_menu(menu_json)

    # 3. Construir .docx
    return _build_docx(menu_json, historial_paciente, calorias_objetivo)

def generate_shopping_list_json(menu_data: dict, gemini_key: str) -> list:
    """
    Genera la lista de compras en formato JSON estructurado.
    """
    import re
    import json
    import traceback
    from google import genai
    from google.genai import types

    todos_ingredientes = menu_data.get("todos_ingredientes", [])
    print(f"[ShoppingList AI] Input ingredients count: {len(todos_ingredientes)}")
    ingredientes_str = "\n".join(todos_ingredientes)

    # Define schema as a raw dict to ensure absolute compatibility across Pydantic/Python versions
    shopping_schema = {
        "type": "OBJECT",
        "properties": {
            "categories": {
                "type": "ARRAY",
                "description": "Lista de categorías principales de compras.",
                "items": {
                    "type": "OBJECT",
                    "properties": {
                        "category": {
                            "type": "STRING",
                            "description": "Categoría con su emoji respectivo (ej. 🥩 PROTEÍNAS, 🥛 LÁCTEOS Y SUSTITUTOS, 🥦 VERDURAS Y HORTALIZAS, 🍎 FRUTAS FRESCAS, 🍞 CEREALES Y TUBÉRCULOS, 🥜 GRASAS Y SEMILLAS, 🧂 DESPENSA Y CONDIMENTOS, 🍵 BEBIDAS)."
                        },
                        "items": {
                            "type": "ARRAY",
                            "description": "Lista de ingredientes de esta categoría.",
                            "items": {
                                "type": "OBJECT",
                                "properties": {
                                    "icon": {
                                        "type": "STRING",
                                        "description": "Un único emoji del alimento (ej. 🥬, 🍗, 🥛, 🍚)."
                                    },
                                    "name": {
                                        "type": "STRING",
                                        "description": "Nombre del ingrediente (ej. Pechuga de pollo, Brócoli)."
                                    },
                                    "amount": {
                                        "type": "STRING",
                                        "description": "Cantidad consolidada total (ej. 1.2 kg, 8 rebanadas)."
                                    },
                                    "tip": {
                                        "type": "STRING",
                                        "description": "Consejo clínico o de compra práctico."
                                    }
                                },
                                "required": ["icon", "name", "amount", "tip"]
                            }
                        }
                    },
                    "required": ["category", "items"]
                }
            }
        },
        "required": ["categories"]
    }
    
    system_prompt = (
        "Eres un asistente experto en nutrición clínica y compras de supermercado inteligentes. "
        "Tu tarea es analizar los ingredientes semanales del plan alimenticio de un paciente y "
        "devolver una lista consolidada, optimizada y agrupada en formato JSON estricto."
    )
    
    prompt = f"""
Analiza los siguientes ingredientes extraídos del plan de alimentación de 7 días del paciente:

{ingredientes_str}

Instrucciones de consolidación:
1. **Suma matemática precisa**: Agrupa y unifica los ingredientes repetidos. Por ejemplo, si el plan pide espinacas el lunes, miércoles y viernes, súmalas en un único ingrediente "Espinacas" consolidando sus porciones en una unidad de supermercado (ej. '2 manojos' o '500g').
2. **Unidades de compra realistas**: Si la suma de proteínas da '1400g', conviértela a una unidad de compra lógica como '1.4 kg'.
3. **Clasificación estricta**: Clasifica cada ingrediente en la categoría correcta de la lista proporcionada en la descripción de 'category'.
4. **Consejos prácticos (`tip`)**: Agrega recomendaciones de selección fresca o de perfil saludable (ej. 'Elegir jamón bajo en sodio', 'Yogur griego sin endulzantes artificiales').

Genera el JSON usando el esquema definido.
"""
    
    import time

    client = genai.Client(api_key=gemini_key)
    models_to_try = _resolve_model(client)

    for model in models_to_try:
        for attempt in range(3):
            try:
                print(f"[ShoppingList AI] Calling Gemini using model: {model} (attempt {attempt + 1})...")
                response = client.models.generate_content(
                    model=model,
                    contents=[{"role": "user", "parts": [{"text": system_prompt + "\n\n" + prompt}]}],
                    config=types.GenerateContentConfig(
                        temperature=0.1,
                        response_mime_type="application/json",
                        response_schema=shopping_schema
                    ),
                )
                text = response.text.strip()
                print(f"[ShoppingList AI] Response text length: {len(text)}")
                data = json.loads(text)
                return data.get("categories", [])
            except Exception as e:
                print(f"[Gemini JSON Shopping List] Error with model {model} on attempt {attempt + 1}: {e}")
                err_str = str(e).lower()
                
                # If it's a rate limit or transient service error, sleep and retry
                if any(kw in err_str for kw in ["429", "resource_exhausted", "quota", "503", "unavailable", "high demand"]):
                    retry_delay = extract_retry_delay(e)
                    sleep_time = retry_delay + 1.0 if retry_delay > 0 else 5.0 * (2 ** attempt)
                    print(f"[ShoppingList AI] Rate limit hit. Sleeping for {sleep_time:.2f}s before retrying...")
                    time.sleep(sleep_time)
                    continue
                
                traceback.print_exc()
                break
            
    print("[ShoppingList AI] ERROR: All models failed to generate content.")
    return [
        {
            "category": "⚠️ ERROR AL GENERAR",
            "items": [{"icon": "❌", "name": "No se pudo conectar con Gemini", "amount": "-", "tip": "Reintente más tarde"}]
        }
    ]


def parse_menu_document_to_json(menu_url: str, gemini_key: str) -> dict:
    """
    Descarga el PDF/DOCX de la dieta del paciente y usa Gemini para estructurarlo
    en un formato JSON premium con tiempos de comida, recetas, ingredientes y reemplazos.
    """
    import io
    import json
    import requests
    import traceback
    from google import genai
    from google.genai import types
    from pypdf import PdfReader
    from docx import Document

    print(f"[MenuParser AI] Downloading menu from: {menu_url}")
    
    # 1. Download file
    response = requests.get(menu_url, timeout=15)
    response.raise_for_status()
    file_bytes = response.content
    content_type = response.headers.get('Content-Type', '')

    # 2. Extract plain text
    full_text = ""
    if 'officedocument.wordprocessingml.document' in content_type or menu_url.endswith('.docx'):
        print("[MenuParser AI] Parsing DOCX...")
        doc = Document(io.BytesIO(file_bytes))
        full_text = "\n".join([p.text for p in doc.paragraphs] + [c.text for t in doc.tables for r in t.rows for c in r.cells])
    else:
        print("[MenuParser AI] Parsing PDF...")
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            full_text += (page.extract_text() or "") + "\n"

    print(f"[MenuParser AI] Extracted text length: {len(full_text)}")
    
    # 3. Define schema for Gemini
    menu_schema = {
        "type": "OBJECT",
        "properties": {
            "paciente_nombre": {"type": "STRING"},
            "fecha_elaboracion": {"type": "STRING"},
            "calorias_totales": {"type": "INTEGER", "description": "Calorías totales recomendadas en el plan, si se mencionan. 0 si no se mencionan."},
            "macronutrientes": {
                "type": "OBJECT",
                "properties": {
                    "proteinas_g": {"type": "INTEGER"},
                    "carbohidratos_g": {"type": "INTEGER"},
                    "grasas_g": {"type": "INTEGER"}
                },
                "required": ["proteinas_g", "carbohidratos_g", "grasas_g"]
            },
            "tipo_plan": {
                "type": "STRING",
                "description": "Indica si el plan es por días ('semanal') o por opciones de menús ('equivalencias_opciones')."
            },
            "secciones": {
                "type": "ARRAY",
                "description": "Lista de menús (ej. Menú Opción 1, Menú Opción 2, Menú Opción 3) o Días (ej. Lunes / Sábado, Martes / Domingo, etc.)",
                "items": {
                    "type": "OBJECT",
                    "properties": {
                        "nombre": {"type": "STRING", "description": "Nombre de la sección o día. Ej: 'Lunes / Sábado' o 'Menú Opción 1'"},
                        "tiempos_comida": {
                            "type": "ARRAY",
                            "description": "Lista de tiempos de comida dentro de esta sección.",
                            "items": {
                                "type": "OBJECT",
                                "properties": {
                                    "tiempo": {"type": "STRING", "description": "Ej: 'Licuado', 'Desayuno', 'Colación 1', 'Comida', 'Colación 2', 'Cena'"},
                                    "hora_sugerida": {"type": "STRING", "description": "Hora sugerida (ej. '08:30 AM'), dejar vacío si no se especifica."},
                                    "emoji": {"type": "STRING", "description": "Emoji representativo de la comida (ej. 🍳, 🍏, 🥗, ☕, 🌙)"},
                                    "platillo": {"type": "STRING", "description": "Nombre del platillo o preparación principal. Ej: 'Sándwich de Pollo' o 'Licuado Verde'"},
                                    "preparacion": {"type": "STRING", "description": "Instrucciones de preparación o receta detalladas si vienen en el texto. De lo contrario, dejar vacío."},
                                    "termino_busqueda_imagen": {"type": "STRING", "description": "Término de búsqueda simple en inglés hiper-preciso optimizado para fotografía gastronómica mexicana o internacional en Unsplash/Pexels (ej. 'mexican chilaquiles green salsa', 'chicken fajitas skillet', 'oatmeal bowl berries', 'mexican enfrijoladas', 'nopal toast', 'beef picadillo', 'tuna avocado salad'). Si el platillo es mexicano o latinoamericano, DEBES incluir la palabra 'mexican' o términos culinarios auténticos. Máximo 3 a 4 palabras en inglés."},
                                    "ingredientes": {
                                        "type": "ARRAY",
                                        "description": "Lista de ingredientes o alimentos individuales.",
                                        "items": {
                                            "type": "OBJECT",
                                            "properties": {
                                                "nombre": {"type": "STRING", "description": "Nombre del ingrediente (ej. Huevo, Espinaca, Pan integral)"},
                                                "cantidad": {"type": "STRING", "description": "Porción y unidad (ej. '2 piezas', '1 taza', '100g')"},
                                                "grupo": {"type": "STRING", "description": "Grupo de equivalentes si viene especificado. Ej: 'Cereales sin grasa', 'Origen animal bajo aporte graso'"},
                                                "reemplazos": {
                                                    "type": "ARRAY",
                                                    "description": "Lista de alimentos alternativos sugeridos en el texto para este ingrediente.",
                                                    "items": {"type": "STRING"}
                                                },
                                                "peso_cocido_crudo": {
                                                    "type": "STRING",
                                                    "description": "Si el ingrediente cambia notablemente de peso/volumen al cocinarse (como arroz, pasta, avena, pechuga de pollo, res, pescado, etc.), especifica el equivalente aproximado crudo vs. cocido (ej. '100g crudo ≈ 75g cocido' o '1/2 taza cocida ≈ 40g cruda'). De lo contrario, dejar vacío."
                                                }
                                            },
                                            "required": ["nombre", "cantidad"]
                                        }
                                    },
                                    "suplementos": {
                                        "type": "ARRAY",
                                        "description": "Suplementos sugeridos para tomar junto con esta comida, si se mencionan.",
                                        "items": {"type": "STRING"}
                                    }
                                },
                                "required": ["tiempo", "platillo", "ingredientes", "termino_busqueda_imagen"]
                            }
                        }
                    },
                    "required": ["nombre", "tiempos_comida"]
                }
            },
            "recomendaciones_generales": {
                "type": "ARRAY",
                "description": "Lista de recomendaciones generales de hidratación, preparación o hábitos descritas en el documento.",
                "items": {"type": "STRING"}
            }
        },
        "required": ["paciente_nombre", "tipo_plan", "secciones"]
    }

    system_prompt = (
        "Eres un asistente experto en nutrición clínica y estructuración de datos. "
        "Tu tarea es analizar el texto crudo del plan de alimentación de un paciente (que puede ser en formato de "
        "recetas semanales o un cuadro de equivalencias por opciones) y transformarlo en un JSON estructurado "
        "estricto. No alteres porciones, ingredientes ni indicaciones médicas. Extrae cuidadosamente todas las "
        "recetas, ingredientes, porciones, grupos de equivalentes y reemplazos de alimentos."
    )

    prompt = f"""
    Analiza el siguiente texto extraído del plan de alimentación del paciente:
    
    {full_text}
    
    Instrucciones de mapeo:
    1. Identifica el nombre del paciente y la fecha.
    2. Determina el tipo de plan: 'semanal' (si tiene columnas de días como 'Lunes', 'Lunes / Sábado', 'Martes', etc.) o 'equivalencias_opciones' (si describe menús alternativos como 'Menú Opción 1', 'Menú Opción 2').
    3. Agrupa por secciones (días o menús alternativos).
    4. En cada sección, identifica los tiempos de comida en orden cronológico (ej: Licuado, Desayuno, Colación 1, Comida, Colación 2, Cena).
    5. Para cada tiempo de comida, extrae:
       - El platillo principal (ej. 'Sándwich de Pollo').
       - Las instrucciones de preparación / receta (si se detallan en el texto plano).
       - La lista detallada de ingredientes, incluyendo sus cantidades exactas y grupos si se mencionan.
       - Si en el texto se sugieren reemplazos o sustitutos para un ingrediente en particular (ej: 'Pan integral (o tortilla de maíz 1 pieza)'), agrégalos al arreglo 'reemplazos' de ese ingrediente.
    6. Extrae recomendaciones y suplementos si están presentes en la parte final del texto.
    7. Para ingredientes que cambian de volumen o peso al cocinarse (como arroz, pasta, avena, lentejas, frijoles, pechuga de pollo, res, pescado, etc.), calcula la conversión equivalente aproximada crudo vs. cocido y coloca obligatoriamente AMBOS valores separados por el signo "≈" en el campo 'peso_cocido_crudo' (ej: si dice '100g de pechuga de pollo (crudo)', pon '100g crudo ≈ 75g cocido'; si dice '1/2 taza de arroz cocido', pon '1/2 taza cocido ≈ 40g crudo'; si dice '1/2 taza de avena', pon '1/2 taza cruda ≈ 1 taza cocida'). Deja el campo vacío si el alimento no cambia notablemente al cocinarse.
    
    Genera el JSON usando el esquema definido.
    """

    import time

    client = genai.Client(api_key=gemini_key)
    models_to_try = _resolve_model(client)

    last_exception = None
    rate_limit_exception = None
    for model in models_to_try:
        for attempt in range(3):
            try:
                print(f"[MenuParser AI] Calling Gemini with model: {model} (attempt {attempt + 1})...")
                response = client.models.generate_content(
                    model=model,
                    contents=[{"role": "user", "parts": [{"text": system_prompt + "\n\n" + prompt}]}],
                    config=types.GenerateContentConfig(
                        temperature=0.1,
                        response_mime_type="application/json",
                        response_schema=menu_schema
                    ),
                )
                text = response.text.strip()
                print(f"[MenuParser AI] Extracted JSON length: {len(text)}")
                return json.loads(text)
            except Exception as e:
                print(f"[MenuParser AI] Error with model {model} on attempt {attempt + 1}: {e}")
                last_exception = e
                err_str = str(e).lower()
                
                # If it's a rate limit or transient service error, sleep and retry
                # However, if limit is 0, it is a permanent quota limit block, so we should skip retries.
                is_transient = any(kw in err_str for kw in ["429", "resource_exhausted", "quota", "503", "unavailable", "high demand"])
                is_zero_quota = "limit: 0" in err_str or "limit:0" in err_str
                
                if is_transient and not is_zero_quota:
                    rate_limit_exception = e
                    retry_delay = extract_retry_delay(e)
                    sleep_time = retry_delay + 1.0 if retry_delay > 0 else 5.0 * (2 ** attempt)
                    if sleep_time > 10.0:
                        print(f"[MenuParser AI] Gemini Rate Limit delay ({sleep_time:.2f}s) is too long. Failing model {model} to prevent timeouts.")
                        break
                    print(f"[MenuParser AI] Gemini Rate Limit hit. Sleeping for {sleep_time:.2f}s before retrying...")
                    time.sleep(sleep_time)
                    continue
                
                traceback.print_exc()
                break

    exception_to_raise = rate_limit_exception or last_exception
    if exception_to_raise:
        raise ValueError(f"No se pudo extraer el menú digitalizado debido a un error de Gemini: {str(exception_to_raise)}")
    raise ValueError("No se pudo extraer el menú digitalizado con ningún modelo de Gemini.")


MEXICAN_DISH_SEARCH_MAP = {
    "chilaquiles": "mexican green chilaquiles with chicken",
    "enfrijoladas": "mexican enfrijoladas black bean sauce",
    "enmoladas": "mexican chicken mole enmoladas",
    "entomatadas": "mexican entomatadas tomato sauce",
    "molletes": "mexican molletes refried beans cheese toast",
    "nopal": "mexican nopal cactus salad tostada",
    "nopales": "mexican cooked nopal cactus dish",
    "sincronizada": "mexican ham cheese quesadilla panela",
    "sincronizadas": "mexican ham cheese quesadilla panela",
    "quesadilla": "mexican cheese quesadilla panela",
    "quesadillas": "mexican cheese quesadillas panela",
    "fajitas": "chicken beef fajitas skillet mexican",
    "picadillo": "mexican ground beef picadillo vegetables",
    "ceviche": "mexican fish ceviche tostada",
    "salpicon": "shredded beef salad mexican salpicon",
    "salpicón": "shredded beef salad mexican salpicon",
    "tinga": "mexican chicken tinga tostadas",
    "licuado": "healthy smoothie bowl oat fruit",
    "batido": "healthy smoothie bowl oat fruit",
    "avena": "oatmeal bowl berries cinnamon",
    "huevo": "mexican scrambled eggs salsa",
    "huevos": "mexican scrambled eggs salsa",
    "caldo": "chicken soup mexican caldo de pollo",
    "sopa": "vegetable chicken soup mexican",
    "pozole": "mexican pozole soup",
    "tacos": "mexican tacos guisado",
    "tostada": "mexican tostada avocado salsa",
    "tostadas": "mexican tostada avocado salsa",
    "panela": "panela cheese salad mexican",
    "atun": "tuna salad avocado bowl",
    "atún": "tuna salad avocado bowl",
    "pechuga": "grilled chicken breast salad",
    "bistec": "mexican steak fajitas skillet",
    "alambre": "mexican steak alambre skillet bell pepper",
}


def normalize_mexican_query(dish_name: str, search_term: str) -> str:
    """
    Normaliza y enriquece la consulta de búsqueda para asegurar que represente
    de forma hiper-precisa la gastronomía mexicana o clínica.
    """
    combined = f"{dish_name or ''} {search_term or ''}".lower()
    
    # 1. Buscar coincidencias en el diccionario gastronómico mexicano
    for key, mapped_query in MEXICAN_DISH_SEARCH_MAP.items():
        if key in combined:
            return mapped_query

    # 2. Si ya viene un término formateado en inglés por la IA, asegurar contexto si es platillo mexicano
    raw_term = (search_term or dish_name or "healthy food").lower()
    mexican_keywords = ["taco", "tostada", "nopal", "salsa", "chile", "frijol", "panela", "tortilla", "tamal", "guacamole"]
    if any(kw in combined for kw in mexican_keywords) and "mexican" not in raw_term:
        return f"mexican {raw_term}"
        
    return raw_term


def get_fallback_image(query: str) -> str:
    """
    Devuelve una imagen de comida saludable y mexicana de alta calidad en HD desde Unsplash
    cuando no hay API Keys configuradas o falla la red.
    """
    query_lower = (query or "").lower()
    
    # 1. Licuados / Batidos / Smoothies
    if any(kw in query_lower for kw in ["smoothie", "licuado", "batido", "juice", "jugo", "bebida"]):
        return "https://images.unsplash.com/photo-1553530666-ba11a7da3888?q=80&w=800&auto=format&fit=crop"
    
    # 2. Desayunos Mexicanos / Huevos / Avena / Molletes / Toast
    if any(kw in query_lower for kw in ["egg", "huevo", "chilaquiles", "mollete", "toast", "avocado", "aguacate", "desayuno", "breakfast", "pan", "avena", "oatmeal", "hotcake"]):
        return "https://images.unsplash.com/photo-1525351484163-7529414344d8?q=80&w=800&auto=format&fit=crop"
    
    # 3. Ensaladas / Bowls / Nopales / Ceviche / Panela
    if any(kw in query_lower for kw in ["salad", "ensalada", "nopal", "bowl", "ceviche", "panela", "verdura", "vegetal"]):
        return "https://images.unsplash.com/photo-1540420773420-3366772f4999?q=80&w=800&auto=format&fit=crop"
        
    # 4. Yogurt / Snacks / Fruta / Colaciones
    if any(kw in query_lower for kw in ["yogurt", "yogur", "fruit", "fruta", "snack", "colacion", "colación", "berries", "fresa", "manzana"]):
        return "https://images.unsplash.com/photo-1488477181946-6428a0291777?q=80&w=800&auto=format&fit=crop"
        
    # 5. Carnes / Pollo / Tacos / Fajitas / Picadillo / Tinga / Comida fuerte
    if any(kw in query_lower for kw in ["chicken", "pollo", "beef", "carne", "fish", "pescado", "tacos", "fajitas", "picadillo", "tinga", "tostadas", "comida", "lunch", "dinner"]):
        return "https://images.unsplash.com/photo-1565299585323-38d6b0865b47?q=80&w=800&auto=format&fit=crop"
        
    # Default fotografía culinaria general
    return "https://images.unsplash.com/photo-1490645935967-10de6ba17061?q=80&w=800&auto=format&fit=crop"


def fetch_unsplash_image(query: str, access_key: str) -> str | None:
    """Consulta la API oficial de Unsplash buscando fotografías gastronómicas landscape."""
    import urllib.parse
    import requests

    if not access_key or not query:
        return None

    try:
        safe_query = urllib.parse.quote(f"{query} food")
        url = f"https://api.unsplash.com/search/photos?query={safe_query}&per_page=3&orientation=landscape"
        headers = {"Authorization": f"Client-ID {access_key}"}
        
        resp = requests.get(url, headers=headers, timeout=4)
        if resp.status_code == 200:
            data = resp.json()
            results = data.get("results", [])
            if results:
                urls = results[0].get("urls", {})
                return urls.get("regular") or urls.get("small")
    except Exception as e:
        print(f"[MenuParser AI] Error consultando Unsplash API para '{query}': {e}")
        
    return None


def fetch_pexels_image(query: str, pexels_key: str) -> str | None:
    """Consulta la API de Pexels buscando imágenes fotográficas landscape de alimentos."""
    import urllib.parse
    import requests

    if not pexels_key or not query:
        return None

    try:
        safe_query = urllib.parse.quote(query)
        url = f"https://api.pexels.com/v1/search?query={safe_query}&per_page=3&orientation=landscape"
        headers = {"Authorization": pexels_key}
        
        resp = requests.get(url, headers=headers, timeout=4)
        if resp.status_code == 200:
            data = resp.json()
            photos = data.get("photos", [])
            if photos:
                srcs = photos[0].get("src", {})
                return srcs.get("landscape") or srcs.get("large")
    except Exception as e:
        print(f"[MenuParser AI] Error consultando Pexels API para '{query}': {e}")

    return None


def fetch_dish_image_url(dish_name: str, search_term: str = "", unsplash_key: str = "", pexels_key: str = "") -> str:
    """
    Motor Híbrido de Imágenes (Unsplash + Pexels + Fallback Mexicano):
    Obtiene la fotografía culinaria más precisa y estética del platillo.
    """
    final_query = normalize_mexican_query(dish_name, search_term)
    
    # 1. Intentar Unsplash API (máxima calidad en fotografía gastronómica)
    if unsplash_key:
        img_url = fetch_unsplash_image(final_query, unsplash_key)
        if img_url:
            return img_url

    # 2. Intentar Pexels API (segundo nivel de respaldo)
    if pexels_key:
        img_url = fetch_pexels_image(final_query, pexels_key)
        if img_url:
            return img_url

    # 3. Fallback categorizado HD si no hay claves o falla la conexión
    return get_fallback_image(final_query)