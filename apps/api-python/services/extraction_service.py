import re
from typing import Optional
from docx import Document

def detect_format(doc: Document) -> str:
    """
    Detecta si el documento es formato 'equivalencias', 'menu123' o 'semanal'.
    """
    for table in doc.tables:
        if len(table.rows) == 1 and len(table.columns) == 2:
            c1_text = table.rows[0].cells[1].text.strip()
            if re.match(r'MEN[ÚU]\s*[123]', c1_text, re.IGNORECASE):
                return "equivalencias"

    for table in doc.tables:
        if len(table.rows) >= 3:
            header_text = " ".join(c.text.lower() for c in table.rows[0].cells)
            dias = ["lunes", "martes", "miercoles", "miércoles", "jueves",
                    "viernes", "sabado", "sábado", "domingo"]
            if sum(1 for d in dias if d in header_text) >= 2:
                return "semanal"

    for table in doc.tables:
        if len(table.columns) >= 5:
            for row in table.rows[:3]:
                for ci in range(2, min(6, len(row.cells))):
                    if re.search(r'(MEN[^A-Z\s]*\s*[123]|DIA\s*[123])',
                                 row.cells[ci].text, re.IGNORECASE):
                        return "menu123"

    return "menu123"

def is_section_banner(table) -> Optional[str]:
    """
    Tabla de 1 columna con texto de sección (DESAYUNO/COMIDA/CENA).
    Devuelve la sección o None.
    """
    if len(table.rows) != 1 or len(table.columns) != 1:
        return None
    text = table.rows[0].cells[0].text.strip().upper()
    for sec in ["DESAYUNO", "COMIDA", "CENA"]:
        if sec in text and len(text) < 30:
            return sec
    return None

def is_dish_card(table) -> bool:
    """Tabla de 2 columnas, 1 fila, donde C1 empieza con 'MENÚ N'"""
    if len(table.rows) != 1 or len(table.columns) != 2:
        return False
    c1 = table.rows[0].cells[1].text.strip()
    return bool(re.match(r'MEN[ÚU]\s*[123]', c1, re.IGNORECASE))

def find_menu_table(doc: Document):
    fmt = detect_format(doc)
    for table in doc.tables:
        if fmt == "semanal":
            if len(table.rows) >= 3 and len(table.columns) >= 3:
                col0_texts = [table.rows[i].cells[0].text.lower()
                              for i in range(1, len(table.rows))]
                if any(t in " ".join(col0_texts)
                       for t in ["desayuno", "comida", "cena"]):
                    return table, fmt
        else:
            if len(table.columns) >= 5:
                for row in table.rows[:3]:
                    for ci in [2, 3, 4]:
                        if ci < len(row.cells):
                            if re.search(r'(MEN[^A-Z\s]*\s*[123]|DIA\s*[123])',
                                         row.cells[ci].text, re.IGNORECASE):
                                return table, fmt
    return None, fmt

def extract_equivalencias_data(doc: Document) -> dict:
    menus = {
        "Menu 1": {"dias": 2, "comidas": {}, "cols": [2]},
        "Menu 2": {"dias": 2, "comidas": {}, "cols": [3]},
        "Menu 3": {"dias": 3, "comidas": {}, "cols": [4]},
    }
    current_section = None
    all_text_blocks = []

    for ti, table in enumerate(doc.tables):
        sec = is_section_banner(table)
        if sec:
            current_section = sec
            continue

        if is_dish_card(table):
            cell1 = table.rows[0].cells[1]
            paras = cell1.paragraphs
            menu_label = ""
            dish_name = ""
            description = ""
            for pi, para in enumerate(paras):
                text = para.text.strip()
                if pi == 0:
                    menu_label = text
                elif pi == 1:
                    dish_name = text
                elif pi == 2:
                    description = text

            m_num = re.search(r'[123]', menu_label)
            if m_num and current_section:
                label = f"Menu {m_num.group()}"
                sec_key = current_section
                if label in menus:
                    existing = menus[label]["comidas"].get(sec_key, "")
                    menus[label]["comidas"][sec_key] = (
                        existing + f"\n{dish_name}: {description}"
                        if existing else f"{dish_name}: {description}"
                    )
            all_text_blocks.append(f"{dish_name} - {description}")
            continue

        if (len(table.rows) == 1 and len(table.columns) == 1):
            text = table.rows[0].cells[0].text.strip()
            if "COLACI" in text.upper():
                all_text_blocks.append(text)
                for label in menus:
                    sec_key = "COLACION_" + (
                        "MAT" if "MATUTINA" in text.upper() else "VES")
                    menus[label]["comidas"][sec_key] = text
                continue

        if len(table.columns) == 5 and current_section:
            for row in table.rows:
                cells = row.cells
                if len(cells) < 5:
                    continue
                if re.search(r'MEN[ÚU]\s*[123]', cells[2].text, re.IGNORECASE):
                    continue
                if cells[0].text.strip() == "" and cells[1].text.strip() == "":
                    nota = cells[2].text.strip()
                    if nota:
                        all_text_blocks.append(nota)
                    continue

                amount = cells[0].text.strip()
                tipo = cells[1].text.strip()
                for ci, label in [(2, "Menu 1"), (3, "Menu 2"), (4, "Menu 3")]:
                    txt = cells[ci].text.strip()
                    if txt and label in menus:
                        existing = menus[label]["comidas"].get(
                            current_section, "")
                        entry = (f"{amount} equiv {tipo}: {txt}" if amount
                                 else txt)
                        menus[label]["comidas"][current_section] = (
                            existing + f"\n{entry}" if existing else entry
                        )
                        all_text_blocks.append(txt)

    return {"menus": menus, "todos_ingredientes": all_text_blocks}

def extract_semanal(doc: Document) -> dict:
    table, _ = find_menu_table(doc)
    if not table:
        return {"menus": {}, "todos_ingredientes": []}
    header_row = table.rows[0]
    n_cols = len(header_row.cells)
    col_multiplicadores = {}
    menu_counter = 1
    menus = {}
    menu_3_cols = []
    assigned_menu = {}
    for ci in range(1, n_cols):
        header = header_row.cells[ci].text.strip()
        mult = (len([x for x in header.split("/") if x.strip()])
                if "/" in header else 1)
        col_multiplicadores[ci] = mult
    for ci in range(1, n_cols):
        if col_multiplicadores[ci] >= 2:
            label = f"Menu {menu_counter}"
            menus[label] = {"dias": col_multiplicadores[ci],
                            "comidas": {}, "cols": [ci]}
            assigned_menu[ci] = label
            menu_counter += 1
        else:
            menu_3_cols.append(ci)
    for i, ci in enumerate(menu_3_cols):
        label = f"Menu 3{'abc'[i] if i < 3 else str(i)}"
        menus[label] = {"dias": 1, "comidas": {}, "cols": [ci]}
        assigned_menu[ci] = label
    for ri in range(1, len(table.rows)):
        row = table.rows[ri]
        tiempo = row.cells[0].text.strip()
        if not tiempo:
            continue
        for ci in range(1, n_cols):
            if ci >= len(row.cells):
                continue
            menu_label = assigned_menu.get(ci)
            if menu_label and menu_label in menus:
                contenido = row.cells[ci].text.strip()
                if contenido:
                    menus[menu_label]["comidas"][tiempo] = contenido
    todos = [txt for m in menus.values() for txt in m["comidas"].values()]
    return {"menus": menus, "todos_ingredientes": todos}

def extract_menu123(doc: Document) -> dict:
    table, _ = find_menu_table(doc)
    if not table:
        return {"menus": {}, "todos_ingredientes": []}
    fmt_sub = "menu13"
    for row in table.rows[:5]:
        for ci in range(2, min(6, len(row.cells))):
            if re.search(r'\bDIA\s*[123]\b', row.cells[ci].text, re.IGNORECASE):
                fmt_sub = "dia123"
                break
    menus = {
        "Menu 1": {"dias": 2, "comidas": {}, "cols": [2]},
        "Menu 2": {"dias": 2, "comidas": {}, "cols": [3]},
        "Menu 3": {"dias": 3, "comidas": {}, "cols": [4]},
    }
    seccion_actual = "DESAYUNO"
    colaciones_vistas = 0
    prev_was_colacion = False
    for row in table.rows:
        if len(row.cells) < 5:
            continue
        col0 = row.cells[0].text.strip()
        if fmt_sub == "dia123":
            col0_up = col0.upper()
            if "DESAYUNO" in col0_up:
                seccion_actual = "DESAYUNO"
            elif "COMIDA" in col0_up and len(col0_up) < 20:
                seccion_actual = "COMIDA"
            elif "CENA" in col0_up and len(col0_up) < 20:
                seccion_actual = "CENA"
        else:
            if "COLACI" in col0.upper():
                colaciones_vistas += 1
                prev_was_colacion = True
                continue
            if prev_was_colacion:
                prev_was_colacion = False
                if colaciones_vistas == 1:
                    seccion_actual = "COMIDA"
                elif colaciones_vistas >= 2:
                    seccion_actual = "CENA"
        for ci, menu_label in [(2, "Menu 1"), (3, "Menu 2"), (4, "Menu 3")]:
            if ci >= len(row.cells):
                continue
            cell_text = row.cells[ci].text.strip()
            if not cell_text:
                continue
            is_meal_name = bool(re.search(
                r'(MEN[^A-Z\s]*\s*[123]|\bDIA\s*[123]\b)',
                cell_text, re.IGNORECASE))
            if is_meal_name:
                nombre = re.sub(
                    r'^(MEN[^A-Z\s]*\s*[123]|\bDIA\s*[123]\b)\s*', '',
                    cell_text, flags=re.IGNORECASE).strip()
                menus[menu_label]["comidas"][seccion_actual] = (
                    menus[menu_label]["comidas"].get(seccion_actual, "")
                    + "\n" + nombre)
            else:
                menus[menu_label]["comidas"][seccion_actual] = (
                    menus[menu_label]["comidas"].get(seccion_actual, "")
                    + "\n" + cell_text)
    todos = [txt for m in menus.values() for txt in m["comidas"].values()]
    return {"menus": menus, "todos_ingredientes": todos}

def extract_menu_data(doc: Document) -> dict:
    fmt = detect_format(doc)
    if fmt == "equivalencias":
        return extract_equivalencias_data(doc)
    elif fmt == "semanal":
        return extract_semanal(doc)
    else:
        return extract_menu123(doc)
