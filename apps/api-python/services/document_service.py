import re
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.constants import (
    NS_A, NS_W, SECCIONES_CON_IMAGEN_EQ, SECCIONES_CON_IMAGEN,
    TIEMPOS_CON_IMAGEN_SEMANAL, GROUP_COLORS, TIP_COLOR, COL_WIDTHS
)

def _get_group_colors(group_name: str) -> dict:
    upper_name = group_name.upper()
    for key, colors in GROUP_COLORS.items():
        if key in upper_name:
            return colors
    return GROUP_COLORS["OTROS"]

def _set_cell_shading(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    tcPr.append(shd)

def _set_cell_borders(cell, color: str = "CCCCCC", size: int = 4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcPr.append(tcBorders)

def _set_col_width(cell, width_dxa: int):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcPr.append(tcW)

def has_image_in_cell(cell) -> bool:
    for run in cell.paragraphs[0].runs if cell.paragraphs else []:
        if 'graphic' in run._r.xml:
            return True
        for drawing in run._r.findall(f'.//{{{NS_A}}}blip'):
            return True
    return False

def extract_meal_name(cell) -> str:
    from services.extraction_service import is_dish_card
    if is_dish_card(cell._tc): # Only valid for table context, adapted
        pass
    text = cell.text.strip()
    if text:
        parts = text.split('\n')
        name_part = parts[0]
        name_part = re.sub(
            r'^(MEN[^A-Z\s]*\s*[123]|\bDIA\s*[123]\b)\s*', '',
            name_part, flags=re.IGNORECASE).strip()
        return name_part.split(':')[0].strip()
    return ""

def _add_shopping_group_table(doc: Document, group_title: str, rows_data: list):
    colors = _get_group_colors(group_title)
    dark = colors["dark"]
    light = colors["light"]
    n_data = len(rows_data)
    tbl = doc.add_table(rows=2 + n_data, cols=5)

    row0 = tbl.rows[0]
    cell0 = row0.cells[0]
    cell0.merge(row0.cells[4])
    _set_cell_shading(cell0, dark)
    _set_cell_borders(cell0, dark)
    p0 = cell0.paragraphs[0]
    run0 = p0.add_run(group_title)
    run0.bold = True
    run0.font.size = Pt(11)
    run0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    row1 = tbl.rows[1]
    headers = ["", "Alimento", "Cantidad 7 días", "💡 Tip de compra", "⭐ Marca recomendada"]
    for ci, (hdr, w) in enumerate(zip(headers, COL_WIDTHS)):
        cell = row1.cells[ci]
        _set_cell_shading(cell, light)
        _set_cell_borders(cell, dark, 4)
        _set_col_width(cell, w)
        p = cell.paragraphs[0]
        r = p.add_run(hdr)
        r.bold = True
        r.font.size = Pt(9)

    for di, row_data in enumerate(rows_data):
        row = tbl.rows[2 + di]
        row_fill = "FFFFFF" if di % 2 == 0 else light
        for ci, (val, w) in enumerate(zip(row_data[:5], COL_WIDTHS)):
            cell = row.cells[ci]
            fill = TIP_COLOR if ci == 3 else row_fill
            _set_cell_shading(cell, fill)
            _set_cell_borders(cell, "CCCCCC", 2)
            _set_col_width(cell, w)
            p = cell.paragraphs[0]
            r = p.add_run(str(val) if val else "")
            r.font.size = Pt(9)
            if ci == 1:
                r.bold = True

def _parse_shopping_markdown(text: str) -> list:
    groups = []
    current_title = None
    current_rows = []
    for line in text.split("\n"):
        line = line.strip()
        if not line: continue
        if line.startswith("## "):
            if current_title is not None:
                groups.append((current_title, current_rows))
            current_title = line[3:].strip()
            current_rows = []
            continue
        if re.match(r'^\|[\s\-\|:]+\|$', line): continue
        if line.startswith("|"):
            parts = [c.strip() for c in line.split("|") if c.strip()]
            if not parts: continue
            if parts[0].lower() in ("icono", "icon", "emoji"): continue
            while len(parts) < 5: parts.append("")
            current_rows.append(parts[:5])
    if current_title is not None:
        groups.append((current_title, current_rows))
    return groups

def _find_shopping_section_tables(doc: Document) -> tuple:
    start_idx = None
    end_idx = None
    for ti, table in enumerate(doc.tables):
        if len(table.rows) == 1 and len(table.columns) == 1:
            text = table.rows[0].cells[0].text.strip().upper()
            if "LISTA DE COMPRAS" in text:
                start_idx = ti
            elif "RECOMENDACIONES" in text and start_idx is not None:
                end_idx = ti
                break
    return start_idx, end_idx

def replace_shopping_tables(doc: Document, markdown_text: str):
    start_idx, end_idx = _find_shopping_section_tables(doc)
    if start_idx is None:
        doc.add_page_break()
        p = doc.add_paragraph()
        r = p.add_run("🛒  LISTA DE COMPRAS SEMANAL (7 DÍAS)")
        r.bold = True
        r.font.size = Pt(14)
        for group_title, rows_data in _parse_shopping_markdown(markdown_text):
            if rows_data:
                _add_shopping_group_table(doc, group_title, rows_data)
                doc.add_paragraph()
        return

    if end_idx is None:
        end_idx = len(doc.tables)

    tables_to_remove = doc.tables[start_idx + 1:end_idx]
    banner_table_el = doc.tables[start_idx]._tbl
    body = doc.element.body
    recom_table_el = (doc.tables[end_idx]._tbl if end_idx < len(doc.tables) else None)

    elements_to_remove = []
    found_banner = False
    found_recom = False
    for child in list(body):
        if child is banner_table_el:
            found_banner = True
            continue
        if recom_table_el is not None and child is recom_table_el:
            found_recom = True
            break
        if found_banner and not found_recom:
            elements_to_remove.append(child)

    for el in elements_to_remove:
        body.remove(el)

    banner_pos = list(body).index(banner_table_el)
    insert_pos = banner_pos + 1
    for group_title, rows_data in _parse_shopping_markdown(markdown_text):
        if not rows_data: continue
        sep_p = OxmlElement('w:p')
        body.insert(insert_pos, sep_p)
        insert_pos += 1
        _add_shopping_group_table(doc, group_title, rows_data)
        new_tbl_el = doc.tables[-1]._tbl
        body.remove(new_tbl_el)
        body.insert(insert_pos, new_tbl_el)
        insert_pos += 1

# =============================================================================
# MANEJO DE IMAGENES
# =============================================================================

def find_image_slots_equivalencias(doc: Document) -> list:
    from services.extraction_service import is_section_banner, is_dish_card
    slots = []
    current_section = None
    for ti, table in enumerate(doc.tables):
        sec = is_section_banner(table)
        if sec:
            current_section = sec
            continue

        if is_dish_card(table):
            cell0 = table.rows[0].cells[0]
            cell1 = table.rows[0].cells[1]
            if current_section in SECCIONES_CON_IMAGEN_EQ:
                has_img = has_image_in_cell(cell0)
                paras = cell1.paragraphs
                meal_name = ""
                if len(paras) > 1:
                    meal_name = paras[1].text.strip().split(':')[0]
                if meal_name and "sobras" not in meal_name.lower():
                    slots.append((ti, meal_name, has_img, current_section))
    return slots

def insert_image_in_cell_equivalencias(doc: Document, table_idx: int, img_bytes: bytes):
    table = doc.tables[table_idx]
    cell0 = table.rows[0].cells[0]
    cell0.text = ""
    p = cell0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    img_stream = io.BytesIO(img_bytes)
    r.add_picture(img_stream, width=Inches(1.2))

def find_image_slots_semanal(doc: Document) -> list:
    from services.extraction_service import find_menu_table
    slots = []
    table, fmt = find_menu_table(doc)
    if not table or fmt != "semanal":
        return slots

    table_idx = doc.tables.index(table)
    for ri, row in enumerate(table.rows[1:], start=1):
        tiempo_text = row.cells[0].text.strip().lower()
        if not hasattr(row, '_cells'): continue
        tiempo = next((t for t in TIEMPOS_CON_IMAGEN_SEMANAL if t in tiempo_text), None)
        if not tiempo: continue
        for ci in range(1, len(row.cells)):
            cell = row.cells[ci]
            text = cell.text.strip()
            if not text: continue
            has_img = has_image_in_cell(cell)
            meal_name = text.split('\n')[0].strip()
            if "libre" in meal_name.lower() or "sobras" in meal_name.lower() or len(meal_name) < 5:
                has_img = True
            slots.append((table_idx, ri, ci, meal_name, has_img, tiempo.upper()))
    return slots

def find_image_slots_menu123(doc: Document) -> list:
    from services.extraction_service import find_menu_table
    slots = []
    table, fmt = find_menu_table(doc)
    if not table or fmt != "menu123":
        return slots

    table_idx = doc.tables.index(table)
    fmt_sub = "menu13"
    for row in table.rows[:5]:
        for ci in range(2, min(6, len(row.cells))):
            if re.search(r'\bDIA\s*[123]\b', row.cells[ci].text, re.IGNORECASE):
                fmt_sub = "dia123"
                break

    seccion_actual = "DESAYUNO"
    colaciones_vistas = 0
    prev_was_colacion = False

    for ri, row in enumerate(table.rows):
        if len(row.cells) < 5: continue
        col0 = row.cells[0].text.strip()
        if fmt_sub == "dia123":
            col0_up = col0.upper()
            if "DESAYUNO" in col0_up: seccion_actual = "DESAYUNO"
            elif "COMIDA" in col0_up and len(col0_up) < 20: seccion_actual = "COMIDA"
            elif "CENA" in col0_up and len(col0_up) < 20: seccion_actual = "CENA"
        else:
            if "COLACI" in col0.upper():
                colaciones_vistas += 1
                prev_was_colacion = True
                continue
            if prev_was_colacion:
                prev_was_colacion = False
                if colaciones_vistas == 1: seccion_actual = "COMIDA"
                elif colaciones_vistas >= 2: seccion_actual = "CENA"

        if seccion_actual not in SECCIONES_CON_IMAGEN: continue

        for ci in [2, 3, 4]:
            if ci >= len(row.cells): continue
            cell = row.cells[ci]
            meal_name = extract_meal_name(cell)
            if not meal_name: continue
            if "libre" in meal_name.lower() or "sobras" in meal_name.lower() or len(meal_name) < 4:
                continue
            has_img = has_image_in_cell(cell)
            slots.append((table_idx, ri, ci, meal_name, has_img, seccion_actual))
    return slots

def insert_image_in_cell(doc: Document, table_idx: int, row_idx: int, col_idx: int, img_bytes: bytes, fmt: str):
    table = doc.tables[table_idx]
    cell = table.rows[row_idx].cells[col_idx]
    original_text = cell.text.strip()
    cell.text = ""
    p_img = cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_img = p_img.add_run()
    img_stream = io.BytesIO(img_bytes)

    # Size logic based on format
    if fmt == "semanal":
        n_cols = len(table.rows[0].cells)
        col_width = 8.5 / n_cols
        img_width = max(0.8, col_width - 0.2)
        r_img.add_picture(img_stream, width=Inches(img_width))
    else:
        r_img.add_picture(img_stream, width=Inches(1.2))

    p_txt = cell.add_paragraph()
    p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_txt = p_txt.add_run(original_text)
    r_txt.font.size = Pt(8.5) if fmt == "semanal" else Pt(10)
