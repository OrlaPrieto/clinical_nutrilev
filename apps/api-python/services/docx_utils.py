from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTES FIJAS
# ─────────────────────────────────────────────────────────────────────────────

NUTRIOLOGA = "Velvet Anakaren De la Cruz Villegas"

# Grupos permitidos
GRUPOS_PERMITIDOS = ["CER-SF", "CER-CF", "POA-M", "POA-Me", "POA-A", "GRA", "VER", "FRU"]

# Alimentos fijos
ALIMENTOS_BENEFICIAN = [
    "🐟 Omega-3: salmón, sardina, chía, linaza molida.",
    "🌿 Antiinflamatorios: cúrcuma, jengibre, canela, té verde.",
    "🥦 Verduras crucíferas cocidas: brócoli, col, coliflor.",
    "🍓 Frutas: fresa, papaya, piña, plátano.",
    "🥬 Verduras cocidas: calabacita, zanahoria, espinaca.",
]

ALIMENTOS_EVITAR = [
    "❌ Embutidos, carnes procesadas y frituras.",
    "❌ Azúcar refinada, postres y bebidas endulzadas.",
    "❌ Lácteos enteros (preferir deslactosados).",
    "❌ Cafeína y bebidas energéticas.",
    "❌ Alimentos ultraprocesados con aditivos.",
    "❌ Alcohol (especialmente vino y cerveza).",
    "❌ Edulcorantes: sorbitol, manitol, xilitol.",
    "❌ El consumo excesivo de sal.",
]

# ─────────────────────────────────────────────────────────────────────────────
# COLORES DEL TEMA
# ─────────────────────────────────────────────────────────────────────────────
COLORS = {
    "header_purple":     "976CA0",
    "header_red":        "EE0000",
    "header_dark_green": "1B6B3A",
    "header_teal":       "1A7A6E",
    "header_desayuno":   "155724",
    "header_moon":       "7D4E00",

    # Grupos de equivalencias
    "grupo_cereal_num":  "664D00",  "grupo_cereal_bg":  "FFF3CD",
    "grupo_poa_num":     "0C5460",  "grupo_poa_bg":     "D1ECF1",
    "grupo_grasa_num":   "7D4E00",  "grupo_grasa_bg":   "FDEBD0",
    "grupo_verdura_num": "155724",  "grupo_verdura_bg": "D4EDDA",
    "grupo_fruta_num":   "7A4000",  "grupo_fruta_bg":   "FDE8C8",

    # Columnas de menú
    "menu1_bg":   "E8F5F3",  "menu1_text": "1A7A6E",
    "menu2_bg":   "FFF3E8",  "menu2_text": "E8762B",
    "menu3_bg":   "F3EEF9",  "menu3_text": "7B3DB5",

    # Colaciones
    "colacion_bg":     "94E4EF",
    "colacion_border": "1A7A6E",
    "colacion_text":   "555555",

    # Composición corporal
    "comp_verde_bg": "F0FAF4",
    "comp_rojo_bg":  "FDE8E2",
    "comp_azul_bg":  "EBF4FA",
    "comp_gris_bg":  "F7F7F7",

    # Info / objetivos / aviso
    "info_bg":     "F0FAF4",
    "info_border": "52B788",
    "info_text":   "1B6B3A",
    "obj_border":  "52B788",
    "warn_bg":     "FFF9E6",
    "warn_border": "F0AD4E",
    "warn_text":   "0C5460",
    "avoid_bg":    "FFF0F0",
    "avoid_border":"E76F51",
    "neutral_bg":  "F7F7F7",

    # Texto
    "text_dark":   "1A1A1A",
    "text_mid":    "444444",
    "text_light":  "777777",
    "text_muted":  "555555",
    "orange_label":"E8762B",
    "teal_label":  "1A7A6E",
    "white":       "FFFFFF",
}

GROUP_COLORS = {
    "CER-SF": (COLORS["grupo_cereal_num"],  COLORS["grupo_cereal_bg"],  "CEREALES"),
    "CER-CF": (COLORS["grupo_cereal_num"],  COLORS["grupo_cereal_bg"],  "CEREALES"),
    "POA-M":  (COLORS["grupo_poa_num"],     COLORS["grupo_poa_bg"],     "POA"),
    "POA-Me": (COLORS["grupo_poa_num"],     COLORS["grupo_poa_bg"],     "POA"),
    "POA-A":  (COLORS["grupo_poa_num"],     COLORS["grupo_poa_bg"],     "POA"),
    "GRA":    (COLORS["grupo_grasa_num"],   COLORS["grupo_grasa_bg"],   "GRASAS"),
    "VER":    (COLORS["grupo_verdura_num"], COLORS["grupo_verdura_bg"], "VERDURAS"),
    "FRU":    (COLORS["grupo_fruta_num"],   COLORS["grupo_fruta_bg"],   "FRUTAS"),
}

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS DE FORMATO
# ─────────────────────────────────────────────────────────────────────────────

def _format_date(date_val) -> str:
    if not date_val: return ""
    s = str(date_val)[:10]
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y"):
        try: return datetime.strptime(s, fmt).strftime("%d-%m-%Y")
        except ValueError: continue
    return s

def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)

def _set_cell_borders(cell, color: str = "CCCCCC", size: int = 3):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color.upper())
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _set_cell_no_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "FFFFFF")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _set_cell_margins(cell, top=80, left=120, bottom=80, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def _set_cell_valign(cell, align: str = "center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)

def _set_table_width(table, width_dxa: int = 10800):
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

def _apply_col_widths(table, widths: list):
    tbl = table._tbl
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None: tblPr.addnext(tblGrid)
    else: tbl.insert(0, tblGrid)

def _add_run(para, text: str, bold=False, size_pt=8.5, color="1A1A1A", font="Arial", italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = _rgb(color)
    return run

def _add_spacing(doc, before_pt=1.5, after_pt=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before_pt)
    p.paragraph_format.space_after = Pt(after_pt)

def _fmt_val(record: dict, keys: list, unit: str = "") -> str:
    for k in keys:
        v = record.get(k)
        if v is not None and v != "" and v != 0:
            return f"{v} {unit}".strip() if unit else str(v)
    return "—"

def _norm_g(g: str) -> str:
    if not g: return ""
    gn = g.upper().replace(" ", "")
    if "FRUT" in gn: return "FRU"
    if "VERD" in gn: return "VER"
    if "CEREAL" in gn: return "CER-SF"
    if "GRASA" in gn: return "GRA"
    if "POA" in gn: return "POA-M"
    return g

# ─────────────────────────────────────────────────────────────────────────────
# BLOQUES ESTRUCTURALES
# ─────────────────────────────────────────────────────────────────────────────

def _add_section_header(doc, text: str, bg_color: str, text_color: str = "FFFFFF", font_size_pt: float = 12.0, bold: bool = True):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _set_table_width(table, 10800)
    cell = table.cell(0, 0)
    _set_cell_no_borders(cell)
    _set_cell_bg(cell, bg_color)
    _set_cell_margins(cell, top=120, left=220, bottom=120, right=220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = _rgb(text_color)
    _add_spacing(doc)

def _add_label_paragraph(doc, emoji_text: str, label: str, color: str = "E8762B", size_pt: float = 9.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, emoji_text + label, bold=True, size_pt=size_pt, color=color)

def _build_info_table(doc, paciente: dict):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.style   = "Table Grid"
    _set_table_width(table, 10800)
    _apply_col_widths(table, [5300, 5500])
    left, right = table.cell(0, 0), table.cell(0, 1)
    for cell in (left, right):
        _set_cell_bg(cell, COLORS["info_bg"])
        _set_cell_borders(cell, COLORS["info_border"], 3)
        _set_cell_margins(cell, 80, 160, 80, 160)
    fecha_fmt = _format_date(paciente.get("fecha_hoy", "")) or "[Fecha de consulta]"
    def _kv(cell, key, val, first=False):
        p = cell.paragraphs[0] if (first and not cell.paragraphs[0].text) else cell.add_paragraph()
        _add_run(p, key, bold=True, size_pt=8.5, color=COLORS["info_text"])
        _add_run(p, val, bold=False, size_pt=8.5, color=COLORS["text_mid"])
        p.paragraph_format.space_after = Pt(1.5)
    _kv(left,  "👤  Paciente: ", paciente.get("nombre", "[Nombre completo]"), first=True)
    _kv(left,  "📅  Fecha: ", fecha_fmt)
    _kv(right, "🗓️  Próxima cita: ", paciente.get("proxima_cita", "[Día, fecha y hora]"), first=True)
    _kv(right, "👩‍⚕️  Nutrióloga: ", NUTRIOLOGA)
    _add_spacing(doc)

def _build_composicion_section(doc, paciente: dict):
    historial = paciente.get("progreso_historial", [])
    ultimo    = historial[0] if historial else {}
    _add_section_header(doc, "📊  COMPOSICIÓN CORPORAL", bg_color=COLORS["header_dark_green"], font_size_pt=9.8)
    fecha_reg = ultimo.get("date") or ultimo.get("created_at") or ""
    if fecha_reg:
        p_f = doc.add_paragraph()
        p_f.paragraph_format.space_before = Pt(0); p_f.paragraph_format.space_after = Pt(3)
        _add_run(p_f, f"  📅 Último registro: {_format_date(fecha_reg)}", italic=True, size_pt=7.5, color=COLORS["text_light"])
    metrics = [
        [(COLORS["comp_verde_bg"], "Peso", _fmt_val(ultimo, ["peso","weight"], "kg")),
         (COLORS["comp_rojo_bg"],  "Masa Grasa", _fmt_val(ultimo, ["masa_grasa","fat_mass"], "kg")),
         (COLORS["comp_verde_bg"], "Masa Magra", _fmt_val(ultimo, ["masa_magra","lean_mass"], "kg")),
         (COLORS["comp_azul_bg"],  "Músculo", _fmt_val(ultimo, ["musculo","muscle"], "kg"))],
        [(COLORS["comp_azul_bg"],  "Agua", _fmt_val(ultimo, ["agua","water"], "L")),
         (COLORS["comp_verde_bg"], "Proteínas", _fmt_val(ultimo, ["proteinas","protein"], "kg")),
         (COLORS["comp_gris_bg"],  "Minerales", _fmt_val(ultimo, ["minerales","minerals"], "kg")),
         (COLORS["comp_rojo_bg"],  "% Grasa", _fmt_val(ultimo, ["pct_grasa","body_fat"], "%"))],
    ]
    tbl = doc.add_table(rows=2, cols=4); tbl.autofit = False
    _set_table_width(tbl, 10800); _apply_col_widths(tbl, [2700]*4)
    for ri, row_data in enumerate(metrics):
        for ci, (bg, label, value) in enumerate(row_data):
            cell = tbl.cell(ri, ci)
            _set_cell_bg(cell, bg); _set_cell_borders(cell, "CCCCCC", 2); _set_cell_margins(cell, 90, 140, 90, 140)
            p_lbl = cell.paragraphs[0]; p_lbl.paragraph_format.space_after = Pt(1)
            _add_run(p_lbl, label, size_pt=7.5, color=COLORS["text_light"])
            p_val = cell.add_paragraph(); p_val.paragraph_format.space_before = Pt(0)
            _add_run(p_val, value, bold=True, size_pt=11.0, color=COLORS["header_dark_green"] if value != "—" else COLORS["text_light"])
    _add_spacing(doc)

def _build_aviso(doc):
    table = doc.add_table(rows=1, cols=1); table.style = "Table Grid"; _set_table_width(table, 10800)
    cell = table.cell(0, 0); _set_cell_bg(cell, COLORS["warn_bg"]); _set_cell_borders(cell, COLORS["warn_border"], 3)
    _set_cell_margins(cell, 100, 200, 100, 200)
    _add_run(cell.paragraphs[0], "⚠️ Cualquier aumento de alimentos u omisión produce un desequilibrio en el requerimiento actual. En caso de duda o incomodidad con algún platillo o alimento, favor de notificar a su nutrióloga para evaluar el cambio.", size_pt=8.0, color=COLORS["warn_text"])
    _add_spacing(doc)

def _build_objetivos_table(doc, menu_json: dict, paciente: dict):
    objetivo     = menu_json.get("metadata", {}).get("objetivo_clinico", "Mejorar hábitos y composición corporal.")
    advertencias = menu_json.get("metadata", {}).get("advertencias_clinicas", [])
    comidas_dia  = paciente.get("comidas_dia", "5–6 veces al día")
    table = doc.add_table(rows=1, cols=2); table.autofit = False; _set_table_width(table, 10800); _apply_col_widths(table, [5300, 5500])
    left, right = table.cell(0, 0), table.cell(0, 1)
    for cell in (left, right):
        _set_cell_bg(cell, COLORS["info_bg"]); _set_cell_borders(cell, COLORS["obj_border"], 3); _set_cell_margins(cell, 100, 160, 100, 160)
    p_l = left.paragraphs[0]; p_l.paragraph_format.space_after = Pt(4)
    _add_run(p_l, "🎯  OBJETIVOS DEL PLAN", bold=True, size_pt=9.0, color=COLORS["info_text"])
    for obj in ([objetivo] + [a[:90] for a in advertencias[:2]]):
        p2 = left.add_paragraph(); p2.paragraph_format.space_after = Pt(1.5)
        _add_run(p2, obj, size_pt=8.0, color=COLORS["text_mid"])
    p_r = right.paragraphs[0]; p_r.paragraph_format.space_after = Pt(4)
    _add_run(p_r, "🍽️  HÁBITOS DE ALIMENTACIÓN", bold=True, size_pt=9.0, color=COLORS["info_text"])
    for h in ["Dormir 7–8 horas diarias.", "Ejercicio 3–4 veces/semana.", "Evitar el ayuno prolongado.", "Tomar 1.5–2 litros de agua.", f"Comer {comidas_dia} cada 3–4 horas."]:
        p2 = right.add_paragraph(); p2.paragraph_format.space_after = Pt(1.5)
        _add_run(p2, h, size_pt=8.0, color=COLORS["text_mid"])
    _add_spacing(doc)

def _build_alimentos_table(doc):
    table = doc.add_table(rows=1, cols=2); table.autofit = False; _set_table_width(table, 10800); _apply_col_widths(table, [5300, 5500])
    left, right = table.cell(0, 0), table.cell(0, 1)
    _set_cell_bg(left, COLORS["info_bg"]); _set_cell_borders(left, COLORS["info_border"], 3); _set_cell_margins(left, 100, 160, 100, 160)
    _add_run(left.paragraphs[0], "🌸  ALIMENTOS QUE BENEFICIAN", bold=True, size_pt=8.5, color=COLORS["info_text"])
    for item in ALIMENTOS_BENEFICIAN:
        _add_run(left.add_paragraph(), item, size_pt=8.0, color=COLORS["text_mid"])
    _set_cell_bg(right, COLORS["avoid_bg"]); _set_cell_borders(right, COLORS["avoid_border"], 3); _set_cell_margins(right, 100, 160, 100, 160)
    _add_run(right.paragraphs[0], "🚫  ALIMENTOS A EVITAR", bold=True, size_pt=8.5, color=COLORS["avoid_border"])
    for item in ALIMENTOS_EVITAR:
        _add_run(right.add_paragraph(), item, size_pt=8.0, color=COLORS["avoid_border"])
    _add_spacing(doc)

def _get_grupos_ordered(platillos: list) -> list[str]:
    grupos = []
    for pl in platillos:
        for eq in pl.get("equivalencias", []):
            g = eq.get("grupo", "")
            if g in GRUPOS_PERMITIDOS and g not in grupos: grupos.append(g)
    return grupos

def _get_desc(platillo: dict, grupo: str) -> str:
    tn = _norm_g(grupo)
    for eq in platillo.get("equivalencias", []):
        if _norm_g(eq.get("grupo", "")) == tn: return eq.get("descripcion", "—") or "—"
    return "—"

def _get_porciones(platillo: dict, grupo: str):
    tn = _norm_g(grupo)
    for eq in platillo.get("equivalencias", []):
        if _norm_g(eq.get("grupo", "")) == tn: return eq.get("porciones", "")
    return ""

def _build_licuado_table(doc, menus_data: dict):
    _add_label_paragraph(doc, "🥤  ", "LICUADO (para todos los menús)", color=COLORS["orange_label"])
    MK = ["menu_1", "menu_2", "menu_3"]
    MB, MT = [COLORS["menu1_bg"], COLORS["menu2_bg"], COLORS["menu3_bg"]], [COLORS["menu1_text"], COLORS["menu2_text"], COLORS["menu3_text"]]
    licuados = [menus_data.get(mk, {}).get("desayuno", {}).get("licuado", {}) for mk in MK]
    grupos = _get_grupos_ordered(licuados) or ["FRU", "VER"]
    table = doc.add_table(rows=1+len(grupos)+1, cols=5); table.autofit = False; _set_table_width(table, 10800); _apply_col_widths(table, [600, 1600, 2867, 2867, 2867])
    for ci in range(5):
        cell = table.rows[0].cells[ci]
        if ci < 2: _set_cell_bg(cell, COLORS["neutral_bg"]); _set_cell_no_borders(cell)
        else:
            idx = ci-2; _set_cell_bg(cell, MB[idx]); _set_cell_no_borders(cell); _set_cell_margins(cell, 60, 100, 60, 100)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _add_run(p, f"MENÚ {idx+1}", bold=True, size_pt=8.0, color=MT[idx])
    for ri, grupo in enumerate(grupos):
        row = table.rows[ri+1]; num_bg, lbl_bg, lbl_text = GROUP_COLORS.get(grupo, ("444444", "F0F0F0", grupo))
        por = _get_porciones(licuados[0], grupo)
        _set_cell_bg(row.cells[0], num_bg); _set_cell_borders(row.cells[0], num_bg, 3); p = row.cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _add_run(p, str(por), bold=True, color=COLORS["white"])
        _set_cell_bg(row.cells[1], lbl_bg); _set_cell_borders(row.cells[1], num_bg, 3); _add_run(row.cells[1].paragraphs[0], lbl_text, bold=True, color=num_bg)
        for mi, lic in enumerate(licuados):
            cell = row.cells[mi+2]; _set_cell_bg(cell, MB[mi]); _set_cell_borders(cell, MT[mi], 2); _add_run(cell.paragraphs[0], _get_desc(lic, grupo), size_pt=7.5)
    _add_run(table.rows[-1].cells[2].paragraphs[0], f"📝 Licuar con agua. Aprox: {licuados[0].get('kcal_total','')} kcal", italic=True, size_pt=7.0, color=COLORS["text_light"])
    _add_spacing(doc)

def _build_platillo_block(doc, menus_data: dict, tiempo: str, equiv_label: str, emoji: str):
    MK = ["menu_1", "menu_2", "menu_3"]
    MB, MT = [COLORS["menu1_bg"], COLORS["menu2_bg"], COLORS["menu3_bg"]], [COLORS["menu1_text"], COLORS["menu2_text"], COLORS["menu3_text"]]
    platillos = [menus_data.get(mk, {}).get(tiempo, {}).get("platillo_solido", {}) if tiempo=="desayuno" else menus_data.get(mk,{}).get(tiempo,{}) for mk in MK]
    if tiempo == "desayuno": _add_label_paragraph(doc, f"{emoji}  ", "PLATILLO PRINCIPAL", color=COLORS["orange_label"])
    for mi, platillo in enumerate(platillos):
        if not platillo: continue
        table = doc.add_table(rows=1, cols=2); table.autofit = False; _set_table_width(table, 10800); _apply_col_widths(table, [1500, 9300])
        b, c = table.cell(0,0), table.cell(0,1)
        _set_cell_bg(b, MB[mi]); _set_cell_no_borders(b); p_b = b.paragraphs[0]; p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER; _add_run(p_b, f"MENÚ {mi+1}", bold=True, color=MT[mi])
        _set_cell_bg(c, MB[mi]); _set_cell_no_borders(c); _set_cell_margins(c, 80, 140, 80, 100)
        _add_run(c.paragraphs[0], platillo.get("nombre",""), bold=True, size_pt=9.0, color=MT[mi])
        _add_run(c.add_paragraph(), f"{platillo.get('descripcion','')}  •  {platillo.get('kcal_total',0)} kcal", size_pt=7.5)
        _add_spacing(doc, 2, 2)
    _add_label_paragraph(doc, "📋  ", f"EQUIVALENTES DE {equiv_label.upper()}", color=COLORS["teal_label"])
    _build_equivalencias_table(doc, platillos)

def _build_equivalencias_table(doc, platillos: list):
    MB, MT = [COLORS["menu1_bg"], COLORS["menu2_bg"], COLORS["menu3_bg"]], [COLORS["menu1_text"], COLORS["menu2_text"], COLORS["menu3_text"]]
    grupos = _get_grupos_ordered(platillos[:1]) or _get_grupos_ordered(platillos)
    if not grupos: return
    table = doc.add_table(rows=1+len(grupos)+1, cols=5); table.autofit = False; _set_table_width(table, 10800); _apply_col_widths(table, [600, 1600, 2867, 2867, 2867])
    for ci in range(5):
        cell = table.rows[0].cells[ci]
        if ci < 2: _set_cell_bg(cell, COLORS["neutral_bg"]); _set_cell_no_borders(cell)
        else:
            idx = ci-2; _set_cell_bg(cell, MB[idx]); _set_cell_no_borders(cell)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _add_run(p, f"MENÚ {idx+1}", bold=True, color=MT[idx])
    for ri, grupo in enumerate(grupos):
        row = table.rows[ri+1]; num_bg, lbl_bg, lbl_text = GROUP_COLORS.get(grupo, ("444444", "F0F0F0", grupo))
        por = _get_porciones(platillos[0], grupo)
        _set_cell_bg(row.cells[0], num_bg); _set_cell_borders(row.cells[0], num_bg, 3); p = row.cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _add_run(p, str(por), bold=True, color=COLORS["white"])
        _set_cell_bg(row.cells[1], lbl_bg); _set_cell_borders(row.cells[1], num_bg, 3); _add_run(row.cells[1].paragraphs[0], lbl_text, bold=True, color=num_bg)
        for mi, platillo in enumerate(platillos):
            cell = row.cells[mi+2]; _set_cell_bg(cell, MB[mi]); _set_cell_borders(cell, MT[mi], 2); _add_run(cell.paragraphs[0], _get_desc(platillo, grupo), size_pt=7.5)
    _add_run(table.rows[-1].cells[2].paragraphs[0], f"📝 Total aprox: {platillos[0].get('kcal_total','')} kcal", italic=True, size_pt=7.0, color=COLORS["text_light"])
    _add_spacing(doc)

def _build_colacion_banner(doc, titulo: str, descripcion: str):
    table = doc.add_table(rows=1, cols=1); table.style = "Table Grid"; _set_table_width(table, 10800)
    cell = table.cell(0, 0); _set_cell_bg(cell, COLORS["colacion_bg"]); _set_cell_borders(cell, COLORS["colacion_border"], 4); _set_cell_margins(cell, 100, 200, 100, 200)
    p = cell.paragraphs[0]; _add_run(p, titulo, bold=True, color=COLORS["colacion_border"]); _add_run(p, descripcion, color=COLORS["colacion_text"])
    _add_spacing(doc)

def _build_docx(menu_json: dict, paciente: dict, calorias: int) -> bytes:
    doc = Document(); section = doc.sections[0]; section.page_width, section.page_height = Cm(21.59), Cm(27.94)
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Cm(1.5)
    _add_section_header(doc, "\t🌿  PLAN DE NUTRICIÓN PERSONALIZADO  🌿\t", bg_color=COLORS["header_purple"])
    _build_info_table(doc, paciente)
    prox = paciente.get("proxima_cita", "[DÍA]  •  [FECHA]  •  [HORA]")
    _add_section_header(doc, f"📅  PRÓXIMA CITA: {prox}", bg_color=COLORS["header_red"], font_size_pt=9.0)
    _build_composicion_section(doc, paciente); _build_aviso(doc); _build_objetivos_table(doc, menu_json, paciente); _build_alimentos_table(doc)
    _add_section_header(doc, "🗓️  PLAN DE MENÚ — EQUIVALENCIAS", bg_color=COLORS["header_teal"])
    menus_data = menu_json.get("menus", menu_json)
    _add_section_header(doc, "🌅  DESAYUNO", bg_color=COLORS["header_desayuno"], font_size_pt=9.8)
    _build_licuado_table(doc, menus_data); _build_platillo_block(doc, menus_data, "desayuno", "EL PLATILLO PRINCIPAL", "🍳")
    col_mat = menu_json.get("colacion_matutina_global")
    col_mat_desc = col_mat.get("descripcion", "") if col_mat else menus_data.get("menu_1",{}).get("colacion_matutina",{}).get("descripcion","")
    if col_mat_desc: _build_colacion_banner(doc, "🍎  COLACIÓN MATUTINA: ", col_mat_desc)
    _add_section_header(doc, "☀️  COMIDA", bg_color=COLORS["header_teal"], font_size_pt=9.8)
    _build_platillo_block(doc, menus_data, "comida", "LA COMIDA", "☀️")
    col_vesp = menu_json.get("colacion_vespertina_global")
    col_vesp_desc = col_vesp.get("descripcion", "") if col_vesp else menus_data.get("menu_1",{}).get("colacion_vespertina",{}).get("descripcion","")
    if col_vesp_desc: _build_colacion_banner(doc, "🥛  COLACIÓN VESPERTINA: ", col_vesp_desc)
    _add_section_header(doc, "🌙  CENA", bg_color=COLORS["header_moon"], font_size_pt=9.8)
    _build_platillo_block(doc, menus_data, "cena", "LA CENA", "🌙")
    _add_spacing(doc)
    footer = doc.add_table(rows=1, cols=3); footer.autofit = False; _set_table_width(footer, 10800); _apply_col_widths(footer, [3600]*3)
    MB, MT = [COLORS["menu1_bg"], COLORS["menu2_bg"], COLORS["menu3_bg"]], [COLORS["menu1_text"], COLORS["menu2_text"], COLORS["menu3_text"]]
    for ci in range(3):
        cell = footer.cell(0, ci); _set_cell_bg(cell, MB[ci]); _set_cell_borders(cell, MT[ci], 3); p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, f"MENÚ {ci+1}  •  {menus_data.get(f'menu_{ci+1}',{}).get('kcal_dia_total',calorias)} kcal/día", bold=True, size_pt=9.0, color=MT[ci])
    buf = BytesIO(); doc.save(buf); return buf.getvalue()
